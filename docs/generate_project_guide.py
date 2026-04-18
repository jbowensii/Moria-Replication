"""Generate the Moria Replication Project Guide as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    return table


def build_document():
    doc = Document()

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Moria Replication Project Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Complete Reference for Reconstructing a UE4.27 Editor Project\nfrom The Lord of the Rings: Return to Moria')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Version 0.2.0\n{datetime.date.today().strftime("%B %d, %Y")}\n\nJohn Bowens').font.size = Pt(12)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS placeholder
    # ========================================================================
    doc.add_heading('Table of Contents', level=1)
    add_para(doc, 'Chapter 1: What This Project Is and Why It Exists')
    add_para(doc, 'Chapter 2: How the Project Is Organized')
    add_para(doc, 'Chapter 3: Setting Up Your Environment')
    add_para(doc, 'Chapter 4: Phase 1 -- Extracting Game Assets')
    add_para(doc, 'Chapter 5: Phase 2 -- Building the Editor Project')
    add_para(doc, 'Chapter 6: Phase 3 -- Importing Assets')
    add_para(doc, 'Chapter 7: Phase 4 -- Meshes and Animations')
    add_para(doc, 'Chapter 8: Phase 5 -- Blueprints and Level Data')
    add_para(doc, 'Chapter 9: Phase 6 -- Level Reconstruction')
    add_para(doc, 'Chapter 10: What Worked and What Did Not')
    add_para(doc, 'Chapter 11: Current Limitations and Next Steps')
    add_para(doc, 'Appendix A: Complete File and Directory Reference')
    add_para(doc, 'Appendix B: Complete Script Reference')
    add_para(doc, 'Appendix C: Tools Reference')
    add_para(doc, 'Appendix D: Asset Type Breakdown')
    add_para(doc, 'Appendix E: Known Issues and Workarounds')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 1: WHAT THIS PROJECT IS
    # ========================================================================
    add_heading(doc, 'Chapter 1: What This Project Is and Why It Exists')

    add_para(doc, 'The Big Picture', bold=True)
    doc.add_paragraph(
        'The Lord of the Rings: Return to Moria is a survival and building game set in the legendary '
        'Mines of Moria from Tolkien\'s world. Players take on the role of Dwarves returning to reclaim '
        'their ancient homeland. They explore procedurally generated underground caves, gather resources '
        'like stone, iron, and mithril, build structures and fortifications, craft weapons and tools, '
        'brew Dwarven beverages, and fight off Orcs, Trolls, and other creatures lurking in the deep.'
    )
    doc.add_paragraph(
        'The game was built by Free Range Games using Unreal Engine 4.27, with a custom framework '
        'they call FGK (Free-range Games Kit). This framework handles everything from how the game '
        'loads its data to how players interact with the building system.'
    )

    add_para(doc, 'The Modding Problem', bold=True)
    doc.add_paragraph(
        'Like many games, Return to Moria has a community of players who want to modify the game -- '
        'adding new building pieces, new items, new visual styles, and fixing things the developers '
        'haven\'t gotten to yet. This is called "modding."'
    )
    doc.add_paragraph(
        'The challenge with Return to Moria is that the game\'s FGK framework has a very specific '
        'way of loading data. When the game starts up, it reads all of its configuration files -- '
        'these are called DataTables, and they define everything: what items exist, what can be built, '
        'what each building piece looks like, what resources are needed to craft things, and so on. '
        'The framework reads all of these tables once at startup and locks them into a memory cache. '
        'After that initial load, the cache is sealed. Nothing new can be added.'
    )
    doc.add_paragraph(
        'This creates a fundamental problem for modders. The most common modding approach is to '
        'inject changes while the game is running (called "runtime injection"). A companion project '
        'called MoriaAdvancedBuilder does exactly this -- it uses a tool called UE4SS to inject a '
        'C++ DLL into the running game. This works great for modifying existing items (changing '
        'recipes, adjusting stats, tweaking behavior), but it cannot add entirely new entries to '
        'the build menu. By the time the runtime mod loads, the FGK cache is already sealed.'
    )

    add_para(doc, 'The Solution', bold=True)
    doc.add_paragraph(
        'The solution is to package mod content the same way the game packages its own content. '
        'The game ships its data in files called "pak" archives (specifically, a newer format called '
        'IoStore). When the game starts, it loads all pak files it finds in its content directory. '
        'If a mod is packaged as a properly formatted pak file, the game loads it alongside its own '
        'content -- before the FGK cache is built. The new DataTable rows, new meshes, new textures, '
        'and new Blueprints are all treated as if they were part of the original game.'
    )
    doc.add_paragraph(
        'This is already proven technology. Existing Moria mods like "Secrets of Khazad-dum" and '
        '"TobiModsAddons" use this exact approach to add new building pieces that appear in the '
        'build menu. The problem is that creating these pak files requires an Unreal Engine editor '
        'project that matches the game\'s structure -- and the game developers don\'t provide one.'
    )

    add_para(doc, 'What This Project Does', bold=True)
    doc.add_paragraph(
        'This project works backward from the shipped game to reconstruct a working Unreal Engine '
        '4.27 editor project. Think of it like taking a finished, published book and reconstructing '
        'the author\'s original manuscript -- not to copy the book, but so you can write new chapters '
        'that fit seamlessly into the story.'
    )
    doc.add_paragraph(
        'The process involves:'
    )
    add_bullet(doc, 'Extracting all 56,797 asset files from the game\'s pak archives')
    add_bullet(doc, 'Generating a compilable UE4.27 project with all the game\'s C++ source headers')
    add_bullet(doc, 'Importing 26,139 assets (textures, materials, meshes, animations, data tables) into the editor')
    add_bullet(doc, 'Decompiling 4,406 Blueprints (visual scripts) into readable pseudocode')
    add_bullet(doc, 'Reconstructing 86 game rooms from their underlying data')
    doc.add_paragraph(
        'The end result is an editor environment where modders can create new content, test it '
        'visually, and then cook (package) it into the pak format the game expects.'
    )

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 2: HOW THE PROJECT IS ORGANIZED
    # ========================================================================
    add_heading(doc, 'Chapter 2: How the Project Is Organized')

    doc.add_paragraph(
        'The project lives in a single repository with a clear directory structure. Understanding '
        'this structure is essential for working with the project.'
    )

    add_heading(doc, 'Top-Level Directories', level=2)

    add_table(doc,
        ['Directory', 'Purpose', 'In Git?'],
        [
            ['project/', 'The UE4.27 editor project -- this is the main deliverable', 'Yes (via Git LFS for binary assets)'],
            ['scripts/', 'All automation scripts (Python) and generated level scripts', 'Yes'],
            ['decompiled/', '4,406 decompiled Blueprint .kms pseudocode files', 'Yes'],
            ['docs/', 'Design documents, research notes, and this guide', 'Yes'],
            ['tools/', 'Third-party tools (FModel, retoc, UModel, etc.)', 'No (gitignored, install locally)'],
        ]
    )

    add_heading(doc, 'The Project Directory (project/)', level=2)
    doc.add_paragraph(
        'This is the heart of the repository -- a fully compilable UE4.27 project.'
    )
    add_table(doc,
        ['Path', 'Contents'],
        [
            ['project/Moria.uproject', 'The project file -- defines 9 game modules and 16 plugins'],
            ['project/Source/', '8,700+ C++ source files across 30 module directories (9 compiled, rest for reference)'],
            ['project/Content/', '26,139 imported assets: textures, materials, meshes, animations, data tables, Blueprints, level maps'],
            ['project/Config/', 'Editor configuration files'],
            ['project/Plugins/', 'Plugins: JsonAsAsset, FGK, EasySkyV2, Voxel, and custom plugins'],
            ['project/Binaries/', 'Compiled output (gitignored, regenerated by build)'],
            ['project/Intermediate/', 'Build intermediates (gitignored)'],
            ['project/Saved/', 'Editor session data (gitignored)'],
        ]
    )

    add_heading(doc, 'The Scripts Directory (scripts/)', level=2)
    doc.add_paragraph(
        'Contains all Python automation scripts organized by pipeline phase, plus 86 generated '
        'level reconstruction scripts in the reconstructed/ subdirectory.'
    )

    add_heading(doc, 'The Decompiled Directory (decompiled/)', level=2)
    doc.add_paragraph(
        'Contains .kms files -- these are Blueprint visual scripts decompiled into a C#-like '
        'pseudocode format by KismetKompiler. The directory structure mirrors the game\'s Content '
        'directory. These files are read-only reference material; they cannot be recompiled back '
        'into working Blueprints.'
    )

    add_heading(doc, 'The Tools Directory (tools/) -- Not in Git', level=2)
    doc.add_paragraph(
        'This directory is gitignored because it contains large third-party executables. Each '
        'contributor must install these tools locally. See Appendix C for the complete list.'
    )
    add_table(doc,
        ['Subdirectory', 'Tool'],
        [
            ['tools/cloud-exports/', '27,313 FModel JSON exports of game assets'],
            ['tools/extracted-assets/', '56,797 raw .uasset extractions from game paks'],
            ['tools/legacy-assets/', '58,549 legacy-format packages converted by retoc'],
            ['tools/retoc/', 'retoc.exe -- IoStore to legacy converter'],
            ['tools/KismetKompiler/', 'KismetKompiler.exe -- Blueprint decompiler'],
            ['tools/JsonAsAssetServer/', 'Cloud Server (Core.exe) for dependency resolution'],
            ['tools/umodel/', 'UModel -- mesh and animation exporter'],
            ['tools/FModel/', 'FModel -- asset viewer and JSON exporter'],
            ['tools/stove-patched.exe', 'Patched Stove for editing .umap files'],
        ]
    )

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 3: SETTING UP YOUR ENVIRONMENT
    # ========================================================================
    add_heading(doc, 'Chapter 3: Setting Up Your Environment')

    add_heading(doc, 'Prerequisites', level=2)
    add_numbered(doc, 'Windows 10 or 11 (64-bit) -- required for UE4.27 and all tools')
    add_numbered(doc, 'Unreal Engine 4.27 -- install via Epic Games Launcher. Install to C:\\Program Files\\Epic Games\\UE_4.27\\')
    add_numbered(doc, 'Visual Studio 2022 -- with "Game development with C++" workload. Needed to compile the project.')
    add_numbered(doc, 'Python 3.10 or later -- for running all automation scripts. Install python-docx if you want to regenerate this document.')
    add_numbered(doc, 'Blender 3.x or later -- for mesh and animation format conversion (PSK/PSA/glTF to FBX).')
    add_numbered(doc, 'The Lord of the Rings: Return to Moria -- installed via Epic Games Store. The scripts need access to the game\'s pak files.')
    add_numbered(doc, 'Git with Git LFS -- the repository uses Git Large File Storage for binary assets.')

    add_heading(doc, 'Installing Third-Party Tools', level=2)
    doc.add_paragraph(
        'After cloning the repository, you need to populate the tools/ directory. These are not '
        'included in git because they are large binaries.'
    )

    add_heading(doc, 'FModel', level=3)
    add_bullet(doc, 'Download from https://fmodel.app')
    add_bullet(doc, 'Extract to tools/FModel/')
    add_bullet(doc, 'Configure it to point at the game\'s pak directory')
    add_bullet(doc, 'Used to export game assets as JSON files for import')

    add_heading(doc, 'retoc', level=3)
    add_bullet(doc, 'Install via: cargo install retoc --root tools/retoc')
    add_bullet(doc, 'Or download a prebuilt binary')
    add_bullet(doc, 'Converts IoStore format (.ucas/.utoc) to legacy .uasset/.uexp format')
    add_bullet(doc, 'Must point at the full Paks/ directory, not individual .utoc files')

    add_heading(doc, 'KismetKompiler', level=3)
    add_bullet(doc, 'Download v0.4.0-alpha from GitHub releases')
    add_bullet(doc, 'Extract to tools/KismetKompiler/')
    add_bullet(doc, 'Decompiles Blueprint bytecode to .kms pseudocode')

    add_heading(doc, 'UModel', level=3)
    add_bullet(doc, 'Download umodel_64.exe from https://www.gildor.org/en/projects/umodel')
    add_bullet(doc, 'Place in tools/umodel/')
    add_bullet(doc, 'Exports meshes and animations from cooked game paks as glTF/PSK/PSA')

    add_heading(doc, 'JsonAsAsset Plugin + Cloud Server', level=3)
    add_bullet(doc, 'JsonAsAsset plugin is already in project/Plugins/JsonAsAsset/')
    add_bullet(doc, 'Download Cloud Server (Core.exe) from https://github.com/Tectors/Core/releases')
    add_bullet(doc, 'Place in tools/JsonAsAssetServer/')
    add_bullet(doc, 'The Cloud Server must be running during batch imports to resolve asset dependencies')

    add_heading(doc, 'Opening the Editor for the First Time', level=2)
    add_numbered(doc, 'Double-click scripts/Moria Editor.lnk (or copy the shortcut to your Desktop)')
    add_numbered(doc, 'If prompted "Missing modules -- rebuild?" click Yes')
    add_numbered(doc, 'If prompted about new plugins, click Dismiss')
    add_numbered(doc, 'Wait for the editor to compile shaders (first launch takes several minutes)')
    add_numbered(doc, 'Expected warning: DefaultMediaTexture missing -- this is harmless')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 4: PHASE 1 -- EXTRACTING GAME ASSETS
    # ========================================================================
    add_heading(doc, 'Chapter 4: Phase 1 -- Extracting Game Assets')

    doc.add_paragraph(
        'The first step is to extract all of the game\'s assets from its pak archives. The game '
        'ships its content in a format called IoStore, which bundles everything into .ucas and .utoc '
        'container files. We need to unpack these into individual .uasset files that we can work with.'
    )

    add_heading(doc, 'Step 1: Export JSON with FModel', level=2)
    add_numbered(doc, 'Open FModel and point it at the game\'s Paks directory: C:\\Program Files\\Epic Games\\ReturnToMoria\\Moria\\Content\\Paks\\')
    add_numbered(doc, 'In FModel, use the bulk export feature to export all assets as JSON')
    add_numbered(doc, 'Save exports to tools/cloud-exports/ (the scripts expect this path)')
    add_numbered(doc, 'This produces approximately 27,313 JSON files describing every asset in the game')
    doc.add_paragraph(
        'The JSON exports contain the full serialized data for each asset -- texture parameters, '
        'material settings, data table rows, curve values, and more. These are what JsonAsAsset '
        'uses to reconstruct assets in the editor.'
    )

    add_heading(doc, 'Step 2: Extract Raw Assets', level=2)
    doc.add_paragraph('Run the extraction script:')
    add_para(doc, '    python scripts/extract_game_assets.py', style='No Spacing')
    doc.add_paragraph(
        'This uses retoc to unpack the game\'s IoStore containers into individual .uasset files. '
        'The result is approximately 56,797 raw .uasset files in tools/extracted-assets/. '
        'These serve as the source material for all subsequent import steps.'
    )
    doc.add_paragraph('Useful flags:')
    add_bullet(doc, '--dry-run : Shows what would be extracted without doing it')
    add_bullet(doc, '--validate : Checks extracted files against FModel exports')
    add_bullet(doc, '--stats : Shows a breakdown by asset type')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 5: PHASE 2 -- BUILDING THE EDITOR PROJECT
    # ========================================================================
    add_heading(doc, 'Chapter 5: Phase 2 -- Building the Editor Project')

    doc.add_paragraph(
        'This phase creates a compilable UE4.27 project that matches the game\'s module and class '
        'structure. The project needs to compile cleanly so the editor can load and the cooking '
        'pipeline can produce valid pak files.'
    )

    add_heading(doc, 'How It Works', level=2)
    doc.add_paragraph(
        'A tool called UE4GameProjectGenerator reads the game\'s Unreal Header Tool (UHT) dumps -- '
        'these are the C++ class declarations that describe every game class, struct, enum, and '
        'function. The generator creates a Source/ directory with header and cpp files for each class.'
    )
    doc.add_paragraph(
        'However, the raw output does not compile. Return to Moria uses a custom fork of Unreal '
        'Engine with modifications to the Gameplay Ability System (GAS), custom engine classes, '
        'and other changes. The stock UE4.27 engine doesn\'t have these, so extensive fixes are '
        'needed. All fixes are documented in scripts/phase2_fix_compilation.md.'
    )

    add_heading(doc, 'Summary of Fixes Required', level=2)
    add_numbered(doc, 'Remove 35+ engine plugin modules from Source/ compilation (keep only 9 game modules)')
    add_numbered(doc, 'Fix .uproject to reference only UE4.27-compatible plugins')
    add_numbered(doc, 'Fix Target.cs files (rename, fix class names, set module list)')
    add_numbered(doc, 'Add missing Build.cs dependencies (GameplayTasks, AudioMixer)')
    add_numbered(doc, 'Create ~21 GAS header redirects for Moria\'s custom engine fork')
    add_numbered(doc, 'Create stub classes for types that only exist in Moria\'s engine (ACullVolume, ANavMeshLockVolume, etc.)')
    add_numbered(doc, 'Fix TEnumAsByte mismatches in UFUNCTION parameters')
    add_numbered(doc, 'Implement pure virtual functions in abstract classes')
    add_numbered(doc, 'Fix constructor patterns for classes with unexported parent constructors')
    add_numbered(doc, 'Reparent classes whose parents have no exported API')
    add_numbered(doc, 'Fix UHT metadata issues (reserved names, missing enum values, delegate declarations)')
    add_numbered(doc, 'Remove interfaces that don\'t exist in stock UE4.27')
    add_numbered(doc, 'Fix copy prevention for structs with deleted copy operators')
    add_numbered(doc, 'Fix delegate signature mismatches')

    add_heading(doc, 'The 9 Compiled Game Modules', level=2)
    add_table(doc,
        ['Module', 'Files', 'Description'],
        [
            ['Moria', '5,114', 'Main game module -- characters, building, inventory, AI, world generation'],
            ['FGK', '1,713', 'Free-range Games Kit framework -- core systems'],
            ['FGKAnalytics', '8', 'Analytics integration'],
            ['FGKDebugMenu', '16', 'Debug menu system'],
            ['FGKLoadingScreen', '6', 'Loading screen'],
            ['FGKNavPowerPlaceholder', '8', 'Navigation placeholder'],
            ['FGKStaticData', '20', 'Static data tables wrapper (this is the cache that blocks runtime mods)'],
            ['FGKUE5Stubs', '17', 'UE5 API stubs for UE4 compatibility'],
            ['FGKUIToolkit', '38', 'UI toolkit'],
        ]
    )

    add_heading(doc, 'Building', level=2)
    doc.add_paragraph('From command line:')
    add_para(doc, '    "C:\\Program Files\\Epic Games\\UE_4.27\\Engine\\Binaries\\DotNET\\UnrealBuildTool.exe" MoriaEditor Win64 Development -Project="<repo>\\project\\Moria.uproject"', style='No Spacing')
    doc.add_paragraph('Expected result: 94/94 actions, 0 errors.')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 6: PHASE 3 -- IMPORTING ASSETS
    # ========================================================================
    add_heading(doc, 'Chapter 6: Phase 3 -- Importing Assets into the Editor')

    doc.add_paragraph(
        'With the editor project compiling, the next step is to populate it with the game\'s assets. '
        'This is the largest and most complex phase, involving multiple tools and import methods.'
    )

    add_heading(doc, 'The JsonAsAsset Pipeline', level=2)
    doc.add_paragraph(
        'The primary import tool is JsonAsAsset, a UE4 editor plugin that reads the JSON files '
        'exported by FModel and reconstructs them as native .uasset files in the editor. It handles '
        'textures, materials, material instances, data tables, curves, sound data, and many other types.'
    )

    add_heading(doc, 'Step 1: Start the Cloud Server', level=3)
    add_para(doc, '    scripts/start_cloud_server.bat', style='No Spacing')
    doc.add_paragraph(
        'The Cloud Server (Core.exe) runs on localhost:1500 and provides dependency resolution. '
        'When JsonAsAsset imports a material that references a texture, the Cloud Server automatically '
        'provides the texture data so both can be imported together. The server must be running '
        'throughout the import process.'
    )

    add_heading(doc, 'Step 2: Run Batch Import', level=3)
    add_para(doc, '    python scripts/batch_import.py --run', style='No Spacing')
    doc.add_paragraph(
        'This script splits the 27,313 JSON exports into small batches (default 200 per batch), '
        'creates temporary directories with the correct FModel path structure, and runs the UE4 '
        'BatchImport commandlet on each batch. If a batch crashes (common due to garbage collection '
        'threading issues in UE4), the script logs the failure and moves to the next batch.'
    )
    doc.add_paragraph('Key features:')
    add_bullet(doc, 'Automatically skips already-imported assets (checks for existing .uasset files)')
    add_bullet(doc, 'Supports resuming from a specific batch number (--start-batch N)')
    add_bullet(doc, 'Can filter by category (--category Art)')
    add_bullet(doc, 'Shows import status (--status)')

    add_heading(doc, 'Step 3: Tiered Import (Alternative)', level=3)
    doc.add_paragraph(
        'For more control, you can import in tiers based on dependency order:'
    )
    add_bullet(doc, 'Tier 1-2 (textures, simple materials): python scripts/run_tier12_import.py --run')
    add_bullet(doc, 'Tier 4 (complex materials, montages): python scripts/run_tier4_import.py --run')
    add_bullet(doc, 'Use scripts/prep_batch_tiers.py to sort assets into tiers first')

    add_heading(doc, 'Import Results', level=2)
    doc.add_paragraph(
        'The JsonAsAsset pipeline successfully imported approximately 11,379 assets including '
        '2,806 textures, materials, data tables, curves, and sound data. This represents the ceiling '
        'of what JsonAsAsset can handle -- remaining asset types (Blueprints, meshes, animations) '
        'require different import methods.'
    )

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 7: PHASE 4 -- MESHES AND ANIMATIONS
    # ========================================================================
    add_heading(doc, 'Chapter 7: Phase 4 -- Meshes and Animations')

    add_heading(doc, 'Static Mesh Import Pipeline', level=2)
    doc.add_paragraph(
        'Static meshes (the 3D models for buildings, rocks, furniture, etc.) cannot be imported via '
        'JsonAsAsset because the JSON only contains metadata (collision settings, material slots, LOD '
        'info) -- not the actual geometry. The actual mesh data must be exported from the cooked '
        'game files using UModel.'
    )

    add_heading(doc, 'The Three-Phase Pipeline', level=3)
    add_numbered(doc, 'UModel Export: UModel reads the cooked .uasset files and exports each mesh as a glTF file')
    add_numbered(doc, 'Blender Conversion: Blender headless mode batch-converts glTF to FBX (the format UE4 imports)')
    add_numbered(doc, 'UE4 Import: The UE4 editor imports the FBX files via commandlet')

    doc.add_paragraph('Running the pipeline:')
    add_para(doc, '    python scripts/mesh_import.py --phase 1    # Export via UModel', style='No Spacing')
    add_para(doc, '    python scripts/mesh_import.py --phase 2    # Convert glTF -> FBX', style='No Spacing')
    add_para(doc, '    python scripts/mesh_import.py --phase 3    # Import into UE4', style='No Spacing')
    add_para(doc, '    python scripts/mesh_import.py --phase all  # Run all three', style='No Spacing')

    doc.add_paragraph(
        'Result: 2,027 static meshes imported with 100% coverage. Every static mesh in the game '
        'is now available in the editor project.'
    )

    add_heading(doc, 'Skeletal Meshes and Destructible Meshes', level=2)
    doc.add_paragraph(
        'Skeletal meshes (animated character models) and destructible meshes (objects that break '
        'apart) could not be imported via the FBX pipeline due to skeleton and physics data '
        'requirements. Instead, 1,322 of these were copied directly from the cooked game files. '
        'They appear in the editor\'s content browser but with limited editing capability.'
    )

    add_heading(doc, 'Animation Import', level=2)
    doc.add_paragraph(
        'Animations were exported from the game as PSA (Unreal skeletal animation) files, then '
        'converted to FBX via Blender, and imported into UE4 with explicit skeleton mapping.'
    )
    doc.add_paragraph('The animation pipeline:')
    add_numbered(doc, 'Build manifest: python scripts/build_anim_fbx_manifest.py')
    add_numbered(doc, 'Convert PSA to FBX: python scripts/blender_psa_to_fbx.py (batch via blender_batch_psk_to_fbx.py)')
    add_numbered(doc, 'Import into UE4: Run scripts/ue4_anim_fbx_import.py inside the UE4 Python console')

    doc.add_paragraph(
        'Result: 4,018 animations imported successfully out of 5,027 PSA files converted to FBX. '
        '1,005 animations already existed in the project from earlier imports. 7 bow/crossbow '
        'animations failed due to extra prop bones not present in the base Dwarf skeleton.'
    )

    add_heading(doc, 'The 7 Failed Bow Animations', level=3)
    doc.add_paragraph('These animations reference a bow mesh bone that doesn\'t exist in the standard Dwarf skeleton:')
    add_bullet(doc, 'Dwa_Bow_Combat_Aim_Fire_Bow')
    add_bullet(doc, 'Dwa_Bow_Combat_Aim_Loop_Bow')
    add_bullet(doc, 'Dwa_Bow_Combat_HipAim_Fire_Bow')
    add_bullet(doc, 'Dwa_Bow_Draw_Aim_Bow')
    add_bullet(doc, 'Dwa_Bow_Draw_HipAim_Bow')
    add_bullet(doc, 'Dwa_Xbow_Combat_Atk_Fire_BOW')
    add_bullet(doc, 'Dwa_Xbow_Combat_Atk_Reload_Bow')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 8: PHASE 5 -- BLUEPRINTS AND LEVEL DATA
    # ========================================================================
    add_heading(doc, 'Chapter 8: Phase 5 -- Blueprints and Level Data')

    add_heading(doc, 'Blueprint Decompilation', level=2)
    doc.add_paragraph(
        'Blueprints are Unreal Engine\'s visual scripting system. They define game logic -- how '
        'enemies behave, how the building system works, what happens when you interact with objects, '
        'and much more. The game ships Blueprints as compiled bytecode inside cooked .uasset files.'
    )
    doc.add_paragraph(
        'We cannot reconstruct working Blueprints from cooked data (the visual graph information '
        'is stripped during cooking), but we can decompile the bytecode into readable pseudocode. '
        'This is invaluable for understanding how game systems work.'
    )

    add_heading(doc, 'The Decompilation Pipeline', level=3)
    add_numbered(doc, 'retoc converts IoStore-format .uasset files to legacy format (.uasset + .uexp)')
    add_numbered(doc, 'KismetKompiler decompiles the legacy .uasset to .kms pseudocode (C#-like syntax)')
    add_numbered(doc, 'Output is converted from UTF-16 to UTF-8 for readability')

    doc.add_paragraph('Running the pipeline:')
    add_para(doc, '    python scripts/batch_decompile_blueprints.py --run', style='No Spacing')
    add_para(doc, '    python scripts/batch_decompile_blueprints.py --run --skip-retoc  # If legacy files already exist', style='No Spacing')

    doc.add_paragraph(
        'Result: 4,406 Blueprints decompiled with 100% success rate. Output is in the decompiled/ '
        'directory, mirroring the game\'s Content directory structure.'
    )

    add_heading(doc, 'Copying Blueprints to the Editor Project', level=2)
    doc.add_paragraph(
        'To browse Blueprints in the editor\'s content browser (even though they can\'t be edited), '
        'copy the legacy-format files into the project:'
    )
    add_para(doc, '    python scripts/copy_legacy_blueprints_to_project.py --run', style='No Spacing')
    doc.add_paragraph(
        'This copies 4,406 Blueprint .uasset/.uexp pairs. The editor can display their properties '
        'and references, but the visual graph is not available.'
    )

    add_heading(doc, 'Level Map Files', level=2)
    doc.add_paragraph(
        'The game\'s level maps (.umap files) can also be copied into the project:'
    )
    add_para(doc, '    python scripts/copy_legacy_umaps_to_project.py --run', style='No Spacing')
    doc.add_paragraph(
        'These are cooked level files and have limited editability in the editor, but they provide '
        'the structural framework for the game\'s levels.'
    )

    add_heading(doc, 'The 12 Skipped Troll AnimMontages', level=3)
    doc.add_paragraph(
        'JsonAsAsset\'s batch import cannot resolve skeleton references that use .6 object indices '
        '(a quirk of how Moria serializes certain assets). All 12 affected assets are Troll animation '
        'montages referencing Troll_Parent_Skeleton.6. BlendSpace, BlendSpace1D, and AimOffset assets '
        'all imported successfully -- only AnimMontage was affected.'
    )

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 9: PHASE 6 -- LEVEL RECONSTRUCTION
    # ========================================================================
    add_heading(doc, 'Chapter 9: Phase 6 -- Level Reconstruction')

    doc.add_paragraph(
        'Return to Moria\'s world is made up of "bubbles" -- self-contained rooms that are stitched '
        'together procedurally. Each bubble has data describing what static meshes go where, what '
        'construction blocks are available, and what breakable objects are scattered around.'
    )
    doc.add_paragraph(
        'This data is stored in MorBubbleData assets. We export these as JSON via FModel, then use '
        'a Python script to generate UE4 editor scripts that reconstruct each room.'
    )

    add_heading(doc, 'How Reconstruction Works', level=2)
    doc.add_paragraph('The reconstruct_level.py script processes each BubbleData JSON in three phases:')
    add_numbered(doc, 'Phase 1 -- InstancedMeshCatalog: Places all static meshes with their correct position, rotation, scale, and material assignments')
    add_numbered(doc, 'Phase 2 -- ConstructionCatalog: Places construction blocks (buildable objects) using transforms from DecoVolumes. A key discovery was that construction blocks have NO transforms of their own -- the position data comes from matching DecoVolume names.')
    add_numbered(doc, 'Phase 3 -- InstancedBreakableCatalog + DecoVolumes: Places breakable objects (rubble, vases, crates) and decorative volume markers')

    add_heading(doc, 'Running a Reconstruction', level=2)
    doc.add_paragraph('Step 1: Generate the UE4 script from a BubbleData JSON:')
    add_para(doc, '    python scripts/reconstruct_level.py <BD_json_file> --output scripts/reconstructed/output.py', style='No Spacing')
    doc.add_paragraph('Step 2: Run the generated script in the UE4 editor:')
    add_numbered(doc, 'Open the editor')
    add_numbered(doc, 'Go to Edit > Execute Python Script')
    add_numbered(doc, 'Select the generated .py file from scripts/reconstructed/')
    add_numbered(doc, 'Wait for placement to complete (large rooms can take several minutes)')

    add_heading(doc, 'Pre-Generated Scripts', level=2)
    doc.add_paragraph(
        'The scripts/reconstructed/ directory contains 86 pre-generated reconstruction scripts, '
        'one per game bubble. Some notable examples:'
    )
    add_table(doc,
        ['Script', 'Room', 'Actor Count'],
        [
            ['reconstruct_BD_BB_Chapter4_BalrogsWake.py', "Balrog's Wake (largest room)", '31,634'],
            ['reconstruct_BD_BB_Chapter4_DurinsForge.py', "Durin's Forge", '~5,000'],
            ['reconstruct_BD_BB_Passage_CrampedRavine.py', 'Cramped Ravine (small, good for testing)', '~200'],
            ['reconstruct_BD_BB_Chapter2_DoorsOfDurin.py', 'Doors of Durin (game start area)', '~3,000'],
            ['reconstruct_BD_BB_Chapter5_MithrilForge.py', 'Mithril Forge', '~4,000'],
        ]
    )
    doc.add_paragraph(
        'Total across all 86 bubbles: 292,543 actors.'
    )

    add_heading(doc, 'What Reconstruction Does Not Include', level=2)
    add_bullet(doc, 'Procedural cave geometry -- the actual cave walls and ceilings are generated at runtime by the game\'s voxel system and are not stored in BubbleData')
    add_bullet(doc, 'Materials and textures -- meshes appear as gray/white because material parent shaders cannot be reconstructed from cooked data')
    add_bullet(doc, 'Lighting -- no lights are placed by the reconstruction scripts; you must add your own')
    add_bullet(doc, 'VFX, particles, fog, and post-processing -- these are handled by Blueprints at runtime')

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 10: WHAT WORKED AND WHAT DID NOT
    # ========================================================================
    add_heading(doc, 'Chapter 10: What Worked and What Did Not')

    add_heading(doc, 'What Worked Well', level=2)

    add_para(doc, 'JsonAsAsset for Data-Driven Assets', bold=True)
    doc.add_paragraph(
        'The JsonAsAsset plugin combined with the Cloud Server proved extremely effective for importing '
        'textures, materials, data tables, curves, and other data-driven assets. The dependency '
        'resolution system handled complex material chains automatically. The batch import script\'s '
        'ability to resume after crashes made the process reliable despite UE4\'s instability with '
        'large import operations.'
    )

    add_para(doc, 'UModel for Mesh Export', bold=True)
    doc.add_paragraph(
        'UModel\'s glTF export produced clean mesh data that converted well through Blender to FBX. '
        'We achieved 100% mesh coverage (2,027/2,027 static meshes). The three-phase pipeline '
        '(UModel > Blender > UE4) was reliable and repeatable.'
    )

    add_para(doc, 'KismetKompiler for Blueprint Decompilation', bold=True)
    doc.add_paragraph(
        'KismetKompiler achieved a 100% success rate decompiling 4,406 Blueprints. The .kms output '
        'is readable and useful for understanding game logic. The retoc + KismetKompiler pipeline '
        'was straightforward once the correct tool versions were identified.'
    )

    add_para(doc, 'BubbleData Reconstruction', bold=True)
    doc.add_paragraph(
        'The level reconstruction approach worked well for placing geometry. The key breakthrough '
        'was discovering that ConstructionCatalog blocks have no transforms of their own -- the '
        'position data comes from DecoVolumes with matching names. This achieved 100% name match '
        'coverage. The 292,543 actors across 86 rooms were placed correctly.'
    )

    add_para(doc, 'retoc for Format Conversion', bold=True)
    doc.add_paragraph(
        'retoc reliably converted all 58,549 IoStore-format packages to legacy format. The key '
        'learning was that it must point at the full Paks/ directory, not individual .utoc files. '
        'The legacy assets cache is reusable across multiple runs (--skip-retoc flag).'
    )

    add_para(doc, 'PowerShell over os.walk for OneDrive', bold=True)
    doc.add_paragraph(
        'A critical discovery: Python\'s os.walk() on a OneDrive-synced directory takes minutes '
        'to complete because it triggers cloud file enumeration. PowerShell\'s Get-ChildItem '
        'completes in seconds. All scripts were updated to use PowerShell for file discovery.'
    )

    add_heading(doc, 'What Did Not Work', level=2)

    add_para(doc, 'Material Parent Shaders', bold=True)
    doc.add_paragraph(
        'This is the single biggest gap. Material Instances were imported with correct parameters '
        '(scalar values, vector values, texture references), but the parent master materials -- '
        'the actual shader graphs that define how textures are combined and rendered -- cannot be '
        'reconstructed from cooked data. The shader graph is compiled into bytecode during cooking '
        'and the original node layout is lost. As a result, all meshes render as gray/white in the '
        'editor. The game uses a complex procedural material system (ProcMaterial_*) that compounds '
        'this problem.'
    )

    add_para(doc, 'Blueprint Reconstruction (not just decompilation)', bold=True)
    doc.add_paragraph(
        'While we can decompile Blueprints to readable pseudocode, we cannot reconstruct working '
        'Blueprint assets. The visual graph data (node positions, connections, pin types) is stripped '
        'during cooking. The cooked .uasset files contain only compiled Kismet bytecode. The '
        'Blueprints copied into the project appear in the content browser but cannot be opened in '
        'the Blueprint editor.'
    )

    add_para(doc, 'Stove (umap Editor) with IoStore Format', bold=True)
    doc.add_paragraph(
        'The Stove tool was patched to handle UE4.27 legacy .umap files (two fixes contributed as '
        'PR #116 to the AstroTechies/unrealmodding repository). However, a third fix for native '
        'IoStore/Zen format was not feasible as a quick patch -- it would require a full Zen parser '
        'or IoStore container reader. The workaround is to use retoc-converted legacy files.'
    )

    add_para(doc, 'OneDrive and Git Interaction', bold=True)
    doc.add_paragraph(
        'OneDrive\'s cloud sync creates "phantom" files that appear in directory listings but don\'t '
        'exist on disk. This caused git add failures and required the --ignore-errors flag. Circular '
        'junction links can self-replicate on OneDrive and must be removed with "cmd /c rmdir". '
        'Working with 21 GB of binary assets on OneDrive added significant friction.'
    )

    add_para(doc, 'Lightmap UV Overlaps', bold=True)
    doc.add_paragraph(
        'Imported meshes have overlapping lightmap UVs, making baked (static) lighting fail. The '
        'game uses dynamic lighting at runtime, so the original meshes were never designed for '
        'lightmap baking. The workaround is to use Movable (dynamic) lights exclusively in the '
        'editor.'
    )

    doc.add_page_break()

    # ========================================================================
    # CHAPTER 11: CURRENT LIMITATIONS AND NEXT STEPS
    # ========================================================================
    add_heading(doc, 'Chapter 11: Current Limitations and Next Steps')

    add_heading(doc, 'Current Limitations', level=2)
    add_numbered(doc, 'No working materials -- meshes are gray/white. Requires creating substitute master materials with texture inputs.')
    add_numbered(doc, 'No playable levels -- there is no game mode, player character, or spawn point configured. Levels are view-only.')
    add_numbered(doc, 'Blueprint logic is read-only -- decompiled pseudocode for reference, not editable graphs.')
    add_numbered(doc, 'Cooking to IoStore paks has not been tested -- the editor compiles but the full cook-and-package workflow is unverified.')
    add_numbered(doc, '7 bow animations and 12 Troll montages are missing (see Known Issues).')

    add_heading(doc, 'Potential Next Steps', level=2)
    add_numbered(doc, 'Create a universal master material with Albedo/Normal/ORM texture inputs to stand in for broken parent shaders')
    add_numbered(doc, 'Write a batch script to reassign all Material Instances to use the substitute master material')
    add_numbered(doc, 'Test the cooking pipeline -- attempt to cook a simple test pak and load it in the game')
    add_numbered(doc, 'Add a test DataTable row (new building piece) and verify it appears in the game\'s build menu')
    add_numbered(doc, 'Investigate Stove Fix 3 for native IoStore/Zen .umap editing')
    add_numbered(doc, 'Explore Blueprint recreation tools for critical game systems')

    doc.add_page_break()

    # ========================================================================
    # APPENDIX A: COMPLETE FILE AND DIRECTORY REFERENCE
    # ========================================================================
    add_heading(doc, 'Appendix A: Complete File and Directory Reference')

    add_heading(doc, 'Repository Root', level=2)
    add_table(doc,
        ['File/Directory', 'Description'],
        [
            ['README.md', 'Project overview, status, quick start guide'],
            ['CLAUDE.md', 'AI assistant context and project memory'],
            ['.gitignore', 'Git exclusion rules (tools/, build artifacts)'],
            ['.gitattributes', 'Git LFS tracking rules for binary assets'],
            ['project/', 'UE4.27 editor project'],
            ['scripts/', 'All automation and reconstruction scripts'],
            ['decompiled/', 'Decompiled Blueprint pseudocode (.kms files)'],
            ['docs/', 'Documentation and research'],
            ['tools/', 'Third-party tools (gitignored)'],
        ]
    )

    add_heading(doc, 'project/Content/ -- Asset Directory Structure', level=2)
    add_table(doc,
        ['Directory', 'Size', 'File Count', 'Contents'],
        [
            ['Art/', '7,019 MB', '6,045', 'Architecture kits, decoration meshes, environmental art, materials'],
            ['CharacterArt/', '6,103 MB', '2,282', 'Character models, armor, creature meshes, textures'],
            ['Unshippable/', '3,026 MB', '1,931', 'Whitebox geometry, third-party assets, test content'],
            ['Maps/', '1,630 MB', '1,640', 'Level map files (.umap)'],
            ['DLC/', '1,004 MB', '1,195', 'DLC content packs (Beorn, Durin\'s Folk, Holiday, etc.)'],
            ['Items/', '821 MB', '1,892', 'Weapons, tools, crafting items, consumables'],
            ['Environments/', '645 MB', '1,189', 'Environmental materials, procedural texturing, lighting'],
            ['FX/', '197 MB', '753', 'Particle effects, VFX assets'],
            ['Animations/', '166 MB', '5,776', 'Animation sequences, montages, blend spaces'],
            ['LevelDesign/', '112 MB', '5,236', 'Level design data, challenges, controllers, deco'],
            ['Tech/', '241 MB', '1,247', 'Technical assets, data tables, configuration'],
            ['UI/', '64 MB', '1,289', 'User interface widgets, icons, textures'],
            ['Character/', '52 MB', '1,262', 'Character Blueprints, AI behaviors, creatures'],
            ['Audio/', '<1 MB', '30', 'Audio metadata (actual audio in Wwise, not in paks)'],
        ]
    )

    add_heading(doc, 'project/Source/ -- Module Structure', level=2)
    doc.add_paragraph(
        'The Source/ directory contains 30 module directories. Only 9 are compiled (listed in '
        '.uproject). The remaining 21 are kept as reference material -- they contain the C++ '
        'headers for engine plugins and third-party modules used by the game, useful for '
        'understanding API surfaces and class hierarchies.'
    )

    doc.add_page_break()

    # ========================================================================
    # APPENDIX B: COMPLETE SCRIPT REFERENCE
    # ========================================================================
    add_heading(doc, 'Appendix B: Complete Script Reference')

    add_heading(doc, 'Phase 1: Extraction', level=2)
    add_table(doc,
        ['Script', 'Purpose', 'Usage'],
        [
            ['extract_game_assets.py', 'Extract raw .uasset files from game paks via retoc', 'python extract_game_assets.py [--dry-run] [--validate] [--stats]'],
        ]
    )

    add_heading(doc, 'Phase 3: Asset Import', level=2)
    add_table(doc,
        ['Script', 'Purpose', 'Usage'],
        [
            ['batch_import.py', 'Main batch importer via JsonAsAsset commandlet', 'python batch_import.py --run [--batch-size N] [--category X] [--start-batch N]'],
            ['bulk_fetch_cloud.py', 'Fetch metadata from Cloud Server', 'python bulk_fetch_cloud.py'],
            ['run_cloud_metadata_pass.py', 'Metadata-only Cloud Server pass', 'python run_cloud_metadata_pass.py'],
            ['run_tier12_import.py', 'Import tier 1-2 assets (textures, simple materials)', 'python run_tier12_import.py --run'],
            ['run_tier4_import.py', 'Import tier 4 assets (complex materials, montages)', 'python run_tier4_import.py --run'],
            ['prep_batch_tiers.py', 'Sort assets into dependency tiers', 'python prep_batch_tiers.py'],
            ['prep_tier4_montage_batch.py', 'Prepare AnimMontage batch for tier 4', 'python prep_tier4_montage_batch.py'],
            ['export_textures.py', 'Export textures for import preparation', 'python export_textures.py'],
            ['ue4_texture_import.py', 'Import textures via UE4 commandlet', 'Run inside UE4 Python console'],
            ['phase3_import.py', 'Earlier import pipeline (superseded by batch_import.py)', 'python phase3_import.py'],
        ]
    )

    add_heading(doc, 'Mesh and Animation Import', level=2)
    add_table(doc,
        ['Script', 'Purpose', 'Usage'],
        [
            ['mesh_import.py', 'Full mesh pipeline: UModel > Blender > UE4', 'python mesh_import.py --phase 1|2|3|all [--dry-run] [--status]'],
            ['gen_mesh_list.py', 'Build list of all meshes to import', 'python gen_mesh_list.py'],
            ['gen_fbx_manifest.py', 'Generate FBX import manifest', 'python gen_fbx_manifest.py'],
            ['ue4_fbx_import.py', 'UE4 commandlet wrapper for FBX import', 'Called by mesh_import.py'],
            ['build_anim_fbx_manifest.py', 'Build animation FBX manifest', 'python build_anim_fbx_manifest.py'],
            ['ue4_anim_fbx_import.py', 'Import animation FBX into UE4', 'Run inside UE4 Python console'],
            ['ue4_fix_bow_anims.py', 'Attempt to fix 7 failed bow animations', 'python ue4_fix_bow_anims.py (currently fails)'],
            ['blender_psk_to_fbx.py', 'Convert single PSK mesh to FBX', 'Called by batch script'],
            ['blender_batch_psk_to_fbx.py', 'Batch convert all PSK to FBX', 'python blender_batch_psk_to_fbx.py'],
            ['blender_gltf_to_fbx.py', 'Convert single glTF to FBX', 'Called by batch script'],
            ['blender_batch_gltf_to_fbx.py', 'Batch convert all glTF to FBX', 'python blender_batch_gltf_to_fbx.py'],
            ['blender_psa_to_fbx.py', 'Convert PSA animations to FBX', 'Called by batch script'],
        ]
    )

    add_heading(doc, 'Blueprint and Level Data', level=2)
    add_table(doc,
        ['Script', 'Purpose', 'Usage'],
        [
            ['batch_decompile_blueprints.py', 'Full decompile pipeline: retoc + KismetKompiler', 'python batch_decompile_blueprints.py --run [--skip-retoc] [--filter NAME]'],
            ['copy_legacy_blueprints_to_project.py', 'Copy legacy Blueprints into editor project', 'python copy_legacy_blueprints_to_project.py --run'],
            ['copy_legacy_umaps_to_project.py', 'Copy legacy .umap files into editor project', 'python copy_legacy_umaps_to_project.py --run'],
            ['reconstruct_level.py', 'Generate room reconstruction script from BubbleData JSON', 'python reconstruct_level.py <BD_json> [--output file.py] [--stats]'],
            ['gap_analysis.py', 'Compare extracted vs imported assets', 'python gap_analysis.py'],
        ]
    )

    add_heading(doc, 'Utility', level=2)
    add_table(doc,
        ['Script', 'Purpose', 'Usage'],
        [
            ['start_cloud_server.bat', 'Launch the JsonAsAsset Cloud Server', 'Double-click or run from terminal'],
            ['Moria Editor.lnk', 'Desktop shortcut to open the UE4 editor', 'Double-click'],
        ]
    )

    doc.add_page_break()

    # ========================================================================
    # APPENDIX C: TOOLS REFERENCE
    # ========================================================================
    add_heading(doc, 'Appendix C: Tools Reference')

    add_table(doc,
        ['Tool', 'Version', 'Purpose', 'Install Location'],
        [
            ['FModel', 'Latest', 'View and export game assets as JSON', 'tools/FModel/'],
            ['retoc', 'v0.1.5+', 'Convert IoStore (.ucas/.utoc) to legacy .uasset/.uexp', 'tools/retoc/'],
            ['KismetKompiler', 'v0.4.0-alpha', 'Decompile Blueprint bytecode to .kms pseudocode', 'tools/KismetKompiler/'],
            ['UModel', 'Latest 64-bit', 'Export meshes and animations from cooked paks', 'tools/umodel/'],
            ['JsonAsAsset', 'UE4.27.2', 'UE4 plugin to import JSON-exported assets', 'project/Plugins/JsonAsAsset/'],
            ['Cloud Server (Core.exe)', 'Latest', 'Dependency resolution server for JsonAsAsset', 'tools/JsonAsAssetServer/'],
            ['Blender', '3.x or 5.x', 'Convert PSK/PSA/glTF to FBX format', 'System install'],
            ['Stove (patched)', 'Custom build', 'Edit .umap level files outside the editor', 'tools/stove-patched.exe'],
            ['UE4.27 Editor', '4.27.2', 'The Unreal Engine editor itself', 'C:\\Program Files\\Epic Games\\UE_4.27\\'],
            ['Visual Studio 2022', 'Latest', 'C++ compiler for project build', 'System install'],
        ]
    )

    doc.add_page_break()

    # ========================================================================
    # APPENDIX D: ASSET TYPE BREAKDOWN
    # ========================================================================
    add_heading(doc, 'Appendix D: Asset Type Breakdown')

    add_heading(doc, 'Import Summary', level=2)
    add_table(doc,
        ['Asset Type', 'Count', 'Import Method', 'Status'],
        [
            ['Textures', '2,806', 'JsonAsAsset + Cloud Server', 'Complete'],
            ['Materials / Material Instances', '~3,000+', 'JsonAsAsset + Cloud Server', 'Imported but non-functional (no parent shaders)'],
            ['Data Tables', '~500+', 'JsonAsAsset', 'Complete'],
            ['Curves / Float Curves', '~200+', 'JsonAsAsset', 'Complete'],
            ['Sound Data / Attenuation', '~100+', 'JsonAsAsset', 'Complete (metadata only, audio in Wwise)'],
            ['Static Meshes', '2,027', 'UModel > Blender > FBX > UE4', 'Complete (100% coverage)'],
            ['Skeletal Meshes', '~800', 'Copied from cooked files', 'View-only in editor'],
            ['Destructible Meshes', '~500', 'Copied from cooked files', 'View-only in editor'],
            ['Animations (FBX)', '4,018', 'PSA > Blender > FBX > UE4', 'Complete (minus 7 bow anims)'],
            ['Animations (cooked)', '4', 'Copied from cooked files', 'Troll AnimSequences'],
            ['Blueprints', '4,406', 'retoc legacy conversion', 'View-only (decompiled to .kms)'],
            ['Level Maps (.umap)', '~200+', 'retoc legacy conversion', 'View-only in editor'],
            ['AnimMontages (Troll)', '12', 'Skipped', 'JsonAsAsset .6 index limitation'],
        ]
    )

    add_heading(doc, 'Total Asset Count', level=2)
    doc.add_paragraph('26,139 assets in the editor project (project/Content/).')
    doc.add_paragraph('56,797 total extracted .uasset files from game paks (tools/extracted-assets/).')
    doc.add_paragraph('27,313 FModel JSON exports (tools/cloud-exports/).')

    doc.add_page_break()

    # ========================================================================
    # APPENDIX E: KNOWN ISSUES AND WORKAROUNDS
    # ========================================================================
    add_heading(doc, 'Appendix E: Known Issues and Workarounds')

    add_heading(doc, 'Materials Render as Gray/White', level=2)
    doc.add_paragraph(
        'Problem: All meshes appear without textures or colors in the editor viewport.'
    )
    doc.add_paragraph(
        'Cause: Material Instances reference parent master materials whose shader graphs were '
        'destroyed during cooking. The editor cannot compile these materials without the original '
        'shader node graph.'
    )
    doc.add_paragraph(
        'Workaround: Switch viewport to Unlit mode to verify meshes are correctly placed. For '
        'visual testing, manually override individual mesh materials with WorldGridMaterial or '
        'create a substitute master material with basic texture inputs.'
    )

    add_heading(doc, 'Lightmap Build Fails with UV Overlap Warnings', level=2)
    doc.add_paragraph(
        'Problem: Building lighting produces hundreds of "overlapping UVs" warnings and black '
        'artifacts.'
    )
    doc.add_paragraph(
        'Cause: Imported meshes use UV channel 0 for lightmaps, but the meshes were designed for '
        'dynamic lighting -- they don\'t have dedicated lightmap UV channels.'
    )
    doc.add_paragraph(
        'Workaround: Set all lights to Movable (dynamic) mobility. Do not use Build Lighting. '
        'Dynamic lights render in real-time without requiring lightmap UVs.'
    )

    add_heading(doc, 'OneDrive Phantom Files', level=2)
    doc.add_paragraph(
        'Problem: git add fails with "No such file or directory" errors for files that appear '
        'to exist in directory listings.'
    )
    doc.add_paragraph(
        'Cause: OneDrive cloud sync creates placeholder entries for cloud-only files that haven\'t '
        'been downloaded locally.'
    )
    doc.add_paragraph(
        'Workaround: Use git add --ignore-errors to skip phantom files. Or right-click the project '
        'folder in Explorer and select "Always keep on this device" to force OneDrive to download '
        'all files.'
    )

    add_heading(doc, 'os.walk() is Extremely Slow on OneDrive', level=2)
    doc.add_paragraph(
        'Problem: Python scripts using os.walk() take minutes to enumerate directories.'
    )
    doc.add_paragraph(
        'Cause: os.walk() triggers OneDrive to check the cloud status of every file.'
    )
    doc.add_paragraph(
        'Workaround: Use PowerShell Get-ChildItem via subprocess instead. All project scripts '
        'have been updated with this pattern.'
    )

    add_heading(doc, 'UE4 Commandlet Crashes During Batch Import', level=2)
    doc.add_paragraph(
        'Problem: The BatchImport commandlet periodically crashes with garbage collection errors.'
    )
    doc.add_paragraph(
        'Cause: UE4.27\'s garbage collector has threading issues when processing large numbers '
        'of assets. The FGK module\'s ensure() calls also trigger early exits.'
    )
    doc.add_paragraph(
        'Workaround: The batch_import.py script handles this automatically -- it catches crashes, '
        'logs the failing batch, and continues with the next batch. Exit code 1 from the commandlet '
        'is normal (FGK ensure error). Parse the log for "BatchImport Complete" to confirm success.'
    )

    add_heading(doc, 'Circular Junction Links on OneDrive', level=2)
    doc.add_paragraph(
        'Problem: Junction links (directory symlinks) can self-replicate in OneDrive-synced '
        'directories, creating infinite recursion.'
    )
    doc.add_paragraph(
        'Workaround: Use "cmd /c rmdir <path>" to remove junction links. Do not use Python\'s '
        'shutil.rmtree() as it will follow the junction and delete the target.'
    )

    # Save
    output_path = r'C:\Users\johnb\OneDrive\Documents\Projects\Moria-Replication\docs\Moria Replication Project Guide.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Sections: 11 chapters + 5 appendices")


if __name__ == '__main__':
    build_document()
