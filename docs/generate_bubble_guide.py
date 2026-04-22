#!/usr/bin/env python3
"""
Generate 'How Bubbles Work in Return to Moria.docx'
A non-technical but detailed guide to the BubbleData system.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x2E, 0x4A, 0x62)

# ── Title ──
title = doc.add_heading('How Bubbles Work in Return to Moria', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Guide to BubbleData, Catalogs, and Object References')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # spacer

# ── Introduction ──
doc.add_heading('What Is a Bubble?', level=1)

doc.add_paragraph(
    'In Return to Moria, every room, hallway, cave, and outdoor area that a player '
    'can explore is called a "bubble." Think of a bubble as a self-contained container '
    'that holds everything needed to build one section of the game world: every pillar, '
    'every pile of rubble, every torch on a wall, every piece of orc scaffolding, and '
    'every breakable rock formation. When the game assembles its procedurally generated '
    'map, it picks from a library of these bubbles and connects them together like '
    'building blocks to create the full underground world of Moria.'
)

doc.add_paragraph(
    'Each bubble is stored as a BubbleData file (a .uasset binary file) that the game '
    'engine reads at startup. These files live in the game\'s data at '
    'Tech/Data/Bubbles/GameWorldCatalog/ and follow a naming convention: they all start '
    'with "BD_BB_" followed by a descriptive name. For example, BD_BB_Chapter2_GameStart '
    'is the starting area (called "Aftermath" in-game), and BD_BB_Chapter4_BalrogsWake '
    'contains both the Balrog\'s Wake and The Desolation rooms.'
)

doc.add_paragraph(
    'A single BubbleData file can actually contain multiple in-game rooms. This happens '
    'because the game uses different "origins" (reference points) within the same file '
    'to define where different rooms begin. For instance, BD_BB_Chapter5_MithrilForge '
    'contains both the Mithril Forge room and the Great Mithril Forge room, each with '
    'its own origin point. There are 87 unique BubbleData files that together define '
    'over 100 distinct in-game locations.'
)

# ── Inside a Bubble ──
doc.add_heading('What Is Inside a BubbleData File?', level=1)

doc.add_paragraph(
    'When you look inside a BubbleData file (after converting it from its binary format '
    'to a readable JSON structure), you find that it is organized into several major '
    'sections called "catalogs." Each catalog is responsible for a different category of '
    'objects. Think of it like a warehouse with different departments: one department '
    'handles decorative scenery, another handles breakable objects, and another handles '
    'things that players have built.'
)

doc.add_paragraph(
    'The three main catalogs are:'
)

# Catalog list
doc.add_heading('The Instanced Mesh Catalog', level=2)

doc.add_paragraph(
    'This is the largest and most important catalog. It contains all of the static '
    'decorative and structural objects in the room: stone pillars, dirt mounds, rubble '
    'piles, banners, rugs, torches, orc fortifications, scaffolding, wooden beams, metal '
    'braces, and thousands of other visual elements that make each room look lived-in '
    'and detailed.'
)

doc.add_paragraph(
    'The Instanced Mesh Catalog uses a system called "instanced meshes" for efficiency. '
    'Instead of storing each individual object separately, it groups identical objects '
    'together into "batches." For example, if a room has 15 identical dirt mounds, '
    'rather than describing each one as a completely separate thing, the catalog creates '
    'one batch entry that says "here is a dirt mound" and then lists 15 different '
    'positions where that dirt mound should appear. This is much more efficient for the '
    'game engine to render.'
)

doc.add_paragraph(
    'Each batch in the Instanced Mesh Catalog contains:'
)

items = [
    ('Mesh Reference', 'A pointer to the 3D model file. This tells the game which '
     'shape to draw. For example, "Dirt_Mound_B" points to a specific dirt mound model.'),
    ('Material References', 'Pointers to the visual surface textures. A single mesh '
     'might reference multiple materials for different parts (the base texture, a detail '
     'overlay, weathering effects, etc.).'),
    ('Collision Profile', 'Defines how the object interacts with players and physics. '
     'Options include "BlockAll" (solid, cannot walk through), "BlockAllDynamic" (solid '
     'but can be affected by physics), "NoCollision" (visual only, players pass through), '
     'and others.'),
    ('Transform List', 'The list of positions where this object appears. Each transform '
     'includes an X, Y, Z position (where it is), rotation (which way it faces), and '
     'scale (how big it is). A single batch might have anywhere from 1 to over 100 '
     'transforms, meaning that many copies of that object exist in the room.'),
    ('Lighting Channels', 'Which lighting system affects this object, controlling how '
     'it responds to torches, ambient light, and other light sources.'),
    ('Custom Data', 'Additional per-instance parameters that can vary between copies, '
     'such as color variation or weathering intensity.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('The Instanced Breakable Catalog', level=2)

doc.add_paragraph(
    'This catalog contains objects that players can destroy: breakable rock formations, '
    'smashable wooden structures, destructible ore deposits, and other objects that '
    'respond to player attacks. When a player swings a pickaxe at a rock wall or smashes '
    'through a wooden barricade, the game is interacting with objects defined in this '
    'catalog.'
)

doc.add_paragraph(
    'The Breakable Catalog is structured differently from the Instanced Mesh Catalog. '
    'Instead of simple mesh-and-position entries, each breakable object has a much richer '
    'definition:'
)

items = [
    ('Breakable Class', 'A reference to a Blueprint class that defines the object\'s '
     'behavior. This class name follows a pattern like "PWM_Quarry_1x1x1_A_C" for '
     'quarry rocks or "Suburbs_Column_Large_A_A_Ruined_C" for ruined columns. The class '
     'determines how the object breaks apart, what resources it drops, and what effects '
     'play when destroyed.'),
    ('Mesh and Materials', 'Similar to the Instanced Mesh Catalog, each breakable has '
     'a visual mesh and materials. However, breakables often have additional mesh '
     'variants for their damaged state.'),
    ('Destruction Properties', 'Settings that control the breaking behavior: how much '
     'force is needed, what impulse to apply to debris, how long debris lingers, what '
     'particle effects and sounds to play.'),
    ('Instances Array', 'The list of positions, but stored differently than the '
     'Instanced Mesh Catalog. Breakable positions are in an "Instances" field rather '
     'than a "Transform" field, and each entry contains the full transform data '
     '(position, rotation, scale).'),
    ('Architecture Properties', 'Flags that describe structural behavior: whether the '
     'object affects navigation (can AI pathfind around it?), whether big creatures '
     'can break through it, and what channels it belongs to.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'An important distinction: the "PWM" prefix on many breakable class names stands '
    'for "Procedural World Mesh" and typically refers to quarry rock formations. These '
    'are the structural rock walls and pillars that form the cave geometry. While they '
    'are technically breakable (players can mine them), they serve a dual purpose as '
    'both resources and structural elements. This is why quarry items need special '
    'handling in any modification system: removing them as a category would strip away '
    'the cave walls themselves.'
)

doc.add_heading('The Construction Catalog', level=2)

doc.add_paragraph(
    'This catalog contains objects that relate to the game\'s building system. '
    'Construction pieces include things like wall sections, floor tiles, and other '
    'elements that either were pre-placed by the game designers or represent the types '
    'of things players can build. In many rooms, the Construction Catalog is relatively '
    'small or empty compared to the other two catalogs.'
)

# ── How Objects Are Referenced ──
doc.add_heading('How Objects Are Referenced', level=1)

doc.add_paragraph(
    'Understanding how objects are identified within a BubbleData file is key to '
    'working with them. Every object has a chain of references that uniquely identifies '
    'what it looks like and how it behaves.'
)

doc.add_heading('The Import Table', level=2)

doc.add_paragraph(
    'At the foundation of every BubbleData file is an "Import Table" -- a master list '
    'of every external resource that the file references. Think of it as a bibliography '
    'at the back of a book: it lists every mesh model, every material texture, every '
    'collision profile, and every Blueprint class that any object in the room needs.'
)

doc.add_paragraph(
    'Each entry in the Import Table has a negative index number (like -1, -2, -3, etc.) '
    'and contains:'
)

items = [
    ('Object Name', 'The human-readable name, like "Dirt_Mound_B" or '
     '"MI_Rubble_Masonry_Pile_03_Inst".'),
    ('Class Name', 'What type of asset it is: "StaticMesh" for 3D models, '
     '"MaterialInstanceConstant" for materials, "Class" for Blueprint classes.'),
    ('Package Path', 'Where the actual file lives in the game\'s file system, like '
     '"/Game/Art/Assets/Kits/Misc/Dirt_Mounds/Dirt_Mound_B".'),
    ('Outer Index', 'A reference to the parent package that contains this asset.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'When a batch in the Instanced Mesh Catalog needs to specify which mesh to use, '
    'it stores a negative number that points into this Import Table. For example, if '
    'Dirt_Mound_B is at position -42 in the Import Table, then any batch that uses '
    'that mesh will have a Mesh value of -42. The same system applies to materials, '
    'collision profiles, and Blueprint classes.'
)

doc.add_heading('Type Rule Strings', level=2)

doc.add_paragraph(
    'For the purposes of identifying and categorizing objects (especially when building '
    'removal rules), each object can be described by a "type rule string" that combines '
    'its key properties into a single, readable identifier. This string follows the '
    'format:'
)

p = doc.add_paragraph()
run = p.add_run('MeshName-Material1-Material2-...-CollisionProfile-Flags')
run.italic = True

doc.add_paragraph(
    'For example, a dirt mound might have the type rule string: '
    '"Dirt_Mound_B-Suburbs_Dirt_Inst-NoCollision-dcl-L100". This tells us that it uses '
    'the Dirt_Mound_B mesh, the Suburbs_Dirt_Inst material, has no collision (you can '
    'walk through it), uses default lighting channels (dcl), and has 100 in its lighting '
    'flags (L100). Two objects with the same type rule string are visually identical -- '
    'they are the same mesh with the same materials, just placed at different positions.'
)

# ── Coordinate System ──
doc.add_heading('The Coordinate System', level=1)

doc.add_paragraph(
    'Every object in a bubble has a position defined by three numbers: X, Y, and Z. '
    'These represent the object\'s location in 3D space, measured in "Unreal Engine '
    'units" (roughly 1 unit = 1 centimeter). The Z axis points upward, while X and Y '
    'define the horizontal plane.'
)

doc.add_heading('World Coordinates vs Local Coordinates', level=2)

doc.add_paragraph(
    'There are two ways to describe where an object is. "World coordinates" describe '
    'the object\'s absolute position in the entire game world. "Local coordinates" '
    'describe the object\'s position relative to the bubble\'s origin point (its center '
    'or reference point). The relationship is simple: local coordinates equal world '
    'coordinates minus the bubble\'s origin. If a bubble\'s origin is at world position '
    '(10000, 20000, 500) and an object is at world position (10300, 20150, 600), then '
    'the object\'s local coordinates are (300, 150, 100).'
)

doc.add_paragraph(
    'Inside BubbleData files, objects are stored using their local coordinates. This '
    'makes sense because the game can place the same bubble at different world positions, '
    'and the internal layout stays the same.'
)

doc.add_heading('Coordinate Transforms and Sign Flips', level=2)

doc.add_paragraph(
    'One complication is that not all BubbleData files store coordinates the same way. '
    'Some files "flip" the X and Y coordinates (negating them) compared to what the '
    'game reports at runtime. For example, if the game says an object is at local '
    'position (-3316, -1574, 212), the BubbleData file might store it as (3316, 1574, '
    '212) with the signs reversed.'
)

doc.add_paragraph(
    'Additionally, some rooms that use generic templates (like the reusable DwarfHall '
    'rooms) can be rotated when placed in the world. A room might be rotated 90 degrees, '
    '180 degrees, or 270 degrees. This means that an object\'s game-reported position '
    'needs to be "unrotated" to match what is stored in the file. For a 90-degree '
    'clockwise rotation, the game coordinates (X, Y) would map to file coordinates '
    '(Y, -X).'
)

doc.add_paragraph(
    'When working with coordinates to identify specific objects, the system tries all '
    'possible transforms (identity, sign flip, and four rotations) and accepts a match '
    'if any transform brings the coordinates within 50 units of each other. This '
    'tolerance accounts for minor floating-point differences between the game\'s runtime '
    'position and the stored position.'
)

# ── Multi-Origin Bubbles ──
doc.add_heading('Multi-Origin Bubbles', level=1)

doc.add_paragraph(
    'As mentioned earlier, a single BubbleData file can contain multiple in-game rooms. '
    'This works through the concept of multiple "origins." Each room within the file has '
    'its own origin point, and objects are associated with a particular room based on '
    'which origin they are closest to or logically grouped with.'
)

doc.add_paragraph(
    'For example, BD_BB_Chapter4_BalrogsWake contains both the Balrog\'s Wake room and '
    'The Desolation room. When you visit Balrog\'s Wake in the game, you see the objects '
    'clustered around one origin point. When you visit The Desolation, you see a '
    'different set of objects clustered around a different origin point, even though they '
    'all live in the same underlying file.'
)

doc.add_paragraph(
    'This has practical implications: any modification that targets objects by type '
    '(like "remove all dirt mounds") will affect both rooms in a multi-origin file. To '
    'target objects in only one of the rooms, you need to use position-based targeting '
    'that specifies the exact coordinates of individual objects.'
)

# ── BubbleData vs BubbleDefs ──
doc.add_heading('BubbleData vs Bubble Definitions', level=1)

doc.add_paragraph(
    'The game uses two separate file systems for bubbles:'
)

items = [
    ('BubbleData (BD_BB_*)', 'These files contain the actual content of each room: '
     'every mesh, material, position, breakable definition, and construction piece. '
     'They are large binary files that define what you see when you enter a room. These '
     'live in Tech/Data/Bubbles/GameWorldCatalog/.'),
    ('Bubble Definitions (BF_BB_*)', 'These files contain metadata about the room: its '
     'name, what zone it belongs to, how it connects to other rooms, what biome it is '
     'associated with, and configuration for the procedural generation system. They are '
     'smaller files that tell the game\'s map generator how and when to use each room. '
     'These live in Tech/Data/BubbleDefs/GameWorldCatalog/.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'Think of it this way: the BubbleData file is the room itself (the furniture, walls, '
    'decorations), while the Bubble Definition is the room\'s entry in the building '
    'directory (its name, floor, which hallways lead to it).'
)

# ── The Bigger Picture ──
doc.add_heading('How It All Fits Together', level=1)

doc.add_paragraph(
    'When the game starts generating a new world, the process works roughly like this:'
)

steps = [
    'The world generation system reads the chapter progression tables to determine what '
    'story milestones need to exist (the Doors of Durin, the Bridge of Khazad-dum, '
    'Dimrill Gate, etc.).',
    'It consults the zone system to determine what types of areas should exist between '
    'those milestones: mining areas, urban districts, cave systems, forges, and so on.',
    'For each zone, it picks from a pool of eligible bubbles using the Bubble Definition '
    'files. Filters and constraints determine which rooms can appear in which zones, and '
    'the zone deck system controls how many of each type appear.',
    'The layout connection system determines how bubbles connect to each other, ensuring '
    'that the map flows logically from one area to the next with appropriate transitions '
    '(passages, highways, vertical shafts, etc.).',
    'Once the layout is determined, the game loads each selected BubbleData file, reads '
    'its catalogs, and instantiates all the objects at their specified positions. The '
    'Instanced Mesh Catalog objects appear as scenery, the Breakable Catalog objects '
    'become mineable resources and destructible structures, and the Construction Catalog '
    'objects become the framework for player building.',
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

# ── Object Removal ──
doc.add_heading('How Object Removal Works (CleanSweep)', level=1)

doc.add_paragraph(
    'The CleanSweep mod removes unwanted clutter objects from BubbleData files before '
    'the game loads them. It operates on the binary .uasset files using a pipeline that '
    'converts them to editable JSON, applies removal rules, and converts them back. The '
    'modified files are packaged into a mod pak that the game loads at startup, '
    'overriding the original files.'
)

doc.add_heading('Type Rules (Global and Bubble-Scoped)', level=2)

doc.add_paragraph(
    'A type rule says "remove every instance of this type of object." When a type rule '
    'targets "Dirt_Mound_B," the system finds every batch in the Instanced Mesh Catalog '
    'whose mesh name contains "Dirt_Mound_B" and removes the entire batch, regardless '
    'of how many individual copies exist. Global type rules apply to every bubble in the '
    'game. Bubble-scoped type rules apply only to a specific named room.'
)

doc.add_heading('Position Entries', level=2)

doc.add_paragraph(
    'A position entry says "remove the specific object at this exact location." This is '
    'used when you want to remove one particular instance of an object without affecting '
    'others of the same type. The system matches by mesh name and local coordinates '
    '(within a tolerance of 50 units), trying all possible coordinate transforms to '
    'account for sign flips and room rotations.'
)

doc.add_heading('Safety Mechanisms', level=2)

doc.add_paragraph(
    'The removal system includes safety mechanisms to prevent crashes. If removing '
    'objects would leave a catalog section completely empty (zero entries in a struct '
    'array), the system keeps one dummy entry. This is necessary because the game\'s '
    'binary format does not handle truly empty arrays gracefully, and the serialization '
    'tool (UAssetGUI) can crash when trying to process them.'
)

# ── Statistics ──
doc.add_heading('By the Numbers', level=1)

stats = [
    ('Total BubbleData files', '87 unique .uasset files'),
    ('Total in-game rooms', 'Over 100 (some files contain multiple rooms)'),
    ('Catalogs per file', '3 (Instanced Mesh, Instanced Breakable, Construction)'),
    ('Largest file', 'BD_BB_Chapter4_BalrogsWake (Balrog\'s Wake + The Desolation) -- '
     'contains over 30,000 object instances'),
    ('Coordinate unit', '1 unit = ~1 centimeter'),
    ('Match tolerance', '50 units (COORD_EPSILON) for position matching'),
    ('CleanSweep coverage', '100 bubble specs, 195 global type rules, 4,872 position '
     'entries, 29,008 objects removed'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for metric, value in stats:
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = value

doc.add_paragraph()  # spacer

# ── Glossary ──
doc.add_heading('Glossary', level=1)

terms = [
    ('Bubble', 'A self-contained room or area in the game, stored as a BubbleData file.'),
    ('BubbleData (BD_BB_*)', 'The binary file containing all objects, meshes, and '
     'positions for a room.'),
    ('Bubble Definition (BF_BB_*)', 'The metadata file describing a room\'s properties '
     'for the world generation system.'),
    ('Catalog', 'A section within a BubbleData file that groups objects by type '
     '(meshes, breakables, constructions).'),
    ('Batch', 'A group of identical objects within a catalog, sharing the same mesh and '
     'materials but placed at different positions.'),
    ('Import Table', 'The master reference list at the end of a BubbleData file, '
     'containing all external assets (meshes, materials, classes) used by the file.'),
    ('Instance', 'A single copy of an object placed at a specific position. One batch '
     'can have many instances.'),
    ('Transform', 'The position (X, Y, Z), rotation, and scale of an object instance.'),
    ('Origin', 'The reference point for a room. Local coordinates are measured relative '
     'to this point.'),
    ('Multi-Origin', 'A BubbleData file that contains objects for multiple in-game '
     'rooms, each with its own origin.'),
    ('Type Rule', 'A removal instruction that targets all instances of a specific mesh '
     'type.'),
    ('Position Entry', 'A removal instruction that targets a specific object at exact '
     'coordinates.'),
    ('IoStore', 'The packaging format used by the game for mod files (.pak, .ucas, '
     '.utoc triplet).'),
    ('PWM', 'Procedural World Mesh -- quarry rocks and structural cave geometry.'),
    ('UAssetGUI', 'Tool that converts binary .uasset files to editable JSON and back.'),
    ('retoc', 'Tool that packages modified .uasset files into IoStore mod paks.'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Term'
hdr_cells[1].text = 'Definition'
for term, defn in terms:
    row_cells = table.add_row().cells
    row_cells[0].text = term
    row_cells[1].text = defn

# ── Save ──
output_path = Path(__file__).parent / 'How Bubbles Work in Return to Moria.docx'
doc.save(str(output_path))
print(f'Saved: {output_path}')
print(f'Pages: ~{len(doc.paragraphs) // 4} (estimated)')
