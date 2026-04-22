#!/usr/bin/env python3
"""
Generate 'How World Generation Works in Return to Moria.docx'
A non-technical but detailed guide to the procedural map generation system.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x2E, 0x4A, 0x62)

# ── Title ──
title = doc.add_heading('How World Generation Works in Return to Moria', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Guide to Procedural Map Layout, Zones, and Probability Controls')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — The Big Picture
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('The Big Picture: How Moria Builds Itself', level=1)

doc.add_paragraph(
    'Every time you start a new game of Return to Moria, the underground world is '
    'assembled fresh from scratch. The mines of Moria are not a single hand-crafted map; '
    'instead, the game uses a layered procedural generation system that picks rooms, '
    'connects them together, and arranges them into a coherent underground world. No two '
    'playthroughs are identical, although they all follow the same story progression and '
    'visit the same landmark locations.'
)

doc.add_paragraph(
    'The generation system is controlled by a set of configuration files (DataTables) '
    'stored in the game\'s data at Tech/Data/GameWorld/. These files define every aspect '
    'of how the world is assembled: what chapters the story follows, what zones exist '
    'within each chapter, what rooms can appear in each zone, how those rooms connect to '
    'each other, and the probability of any given room appearing in a particular '
    'playthrough.'
)

doc.add_paragraph(
    'At the highest level, the system works in four layers, each feeding into the next:'
)

layers = [
    ('Chapters', 'The story progression layer. Chapters define the major milestones of '
     'the game (entering Moria, crossing the Bridge of Khazad-dum, reaching the Dimrill '
     'Gate) and determine what zones are active during each phase of gameplay.'),
    ('Zones', 'The regional layer. Each chapter contains one or more zones, and each zone '
     'represents a themed region of the mines: a mining district, a forge complex, a '
     'flooded cavern system, an orc-occupied settlement, and so on. Zones control how many '
     'rooms to generate and with what probability.'),
    ('Zone Deck (Bubble Selection)', 'The room selection layer. Each zone has a "deck" of '
     'eligible rooms (bubbles) that can appear within it. The deck system uses a card-game '
     'metaphor: rooms are drawn from the deck, and each room has rules about how often it '
     'can appear (guaranteed once, at most once, or repeatable).'),
    ('Layout Graph (Connections)', 'The topology layer. Once rooms have been selected, the '
     'layout system determines how they physically connect to each other. A connection '
     'graph defines which rooms can link to which other rooms, and constraints enforce '
     'that certain critical paths always exist.'),
]

for layer_name, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f'{layer_name}: ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — Chapters
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapters: The Story Skeleton', level=1)

doc.add_paragraph(
    'The game\'s progression is divided into chapters, defined in the file '
    'DT_Moria_Chapters. Each chapter represents a major story phase and determines which '
    'zones are generated during that portion of the game. Think of chapters as the '
    'backbone of the world: they ensure that no matter how randomly the rooms are shuffled, '
    'the player always encounters the key story locations in the right order.'
)

doc.add_paragraph(
    'Each chapter entry contains:'
)

items = [
    ('Zone References', 'A list of which zones belong to this chapter. A chapter might '
     'contain a mix of required story zones and optional side areas.'),
    ('Story Milestones', 'Key locations that must appear in this chapter (like a forge, '
     'a gate, or a boss arena). These are marked as "Required" in the zone deck to '
     'guarantee they always generate.'),
    ('Progression Gates', 'Conditions that must be met before the player can advance to '
     'the next chapter. This ties into the layout constraints system to ensure the map '
     'is always completable.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'The chapter system is defined in DT_Moria_Chapters.uasset. When the world generator '
    'starts building a new map, it reads the chapter progression to know what zones to '
    'create and in what order. Each chapter specifies its zones, and the generator then '
    'hands off to the zone system to populate each region.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — Zones
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Zones: The Regional Architects', level=1)

doc.add_paragraph(
    'Zones are the workhorses of the generation system. Each zone represents a distinct '
    'region of the game world with its own theme, size, and generation rules. Zones are '
    'defined in DT_Moria_Zones and contain several critical parameters that directly '
    'control how the procedural generation behaves.'
)

doc.add_heading('Zone Parameters', level=2)

doc.add_paragraph(
    'Every zone entry has the following key fields:'
)

params = [
    ('TargetBubbles (integer)', 'How many rooms the generator should try to place in this '
     'zone. This is a target, not a guarantee — the actual count depends on available '
     'connections and deck contents. A zone with TargetBubbles of 8 will attempt to place '
     '8 rooms, but might end up with 6 or 10 depending on how the layout resolves.'),
    ('NewBubbleChance (float, 0.0 to 1.0)', 'The probability that the generator will '
     'try to place a new room each time it has an opportunity. A value of 0.8 means there '
     'is an 80% chance of placing another room at each step. Lower values create sparser '
     'zones; higher values create denser ones.'),
    ('AdditionalOpeningChance (float, 0.0 to 1.0)', 'The probability of creating extra '
     'connections between rooms beyond the minimum required. Higher values produce more '
     'interconnected layouts with multiple paths between areas. Lower values produce more '
     'linear, corridor-like layouts.'),
    ('GenerationPriority (integer)', 'Determines the order in which zones are generated. '
     'Higher-priority zones are built first, ensuring that critical story areas are placed '
     'before optional side content. If two zones compete for the same space, the higher-'
     'priority zone wins.'),
]

for title_text, desc in params:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('What This Means for Modding', level=2)

doc.add_paragraph(
    'The zone parameters are the most direct levers for changing how the generated world '
    'feels. If you want a chapter to have more rooms, increase TargetBubbles for its zones. '
    'If you want zones to feel more maze-like with multiple paths, increase '
    'AdditionalOpeningChance. If you want sparser exploration with fewer rooms, decrease '
    'NewBubbleChance. These parameters are all stored in DT_Moria_Zones.uasset and can be '
    'modified through the same UAssetGUI pipeline used for BubbleData.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — Zone Deck
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('The Zone Deck: Choosing Which Rooms Appear', level=1)

doc.add_paragraph(
    'Once the generator knows what zones to create and how big they should be, it needs '
    'to decide which specific rooms (bubbles) to place in each zone. This is controlled '
    'by the Zone Deck system, defined in DT_Moria_ZoneDeck. The metaphor is a card game: '
    'each zone has a deck of cards, and each card represents a room that could be placed '
    'in that zone.'
)

doc.add_heading('The Appearances Enum', level=2)

doc.add_paragraph(
    'The most important field in the zone deck is the "Appearances" enum, which controls '
    'how each room can be used. There are three possible values:'
)

appearances = [
    ('Required', 'This room is guaranteed to appear in every playthrough. It will always '
     'be placed in this zone, no matter what. Story-critical locations like the Bridge of '
     'Khazad-dum, the starting Aftermath area, and the Dimrill Gate are marked as Required '
     'to ensure the game is always completable.'),
    ('Single', 'This room can appear at most once per playthrough. It goes into the deck '
     'as a single card. If it gets drawn during generation, it appears; if not, this '
     'particular playthrough will not have it. This is used for unique optional locations '
     'that should feel special when discovered.'),
    ('Multiple', 'This room can appear more than once. It stays in the deck even after '
     'being drawn, so the generator can place multiple copies. This is used for generic '
     'rooms that provide variety and fill out zones: mining tunnels, hallways, small caves, '
     'and other repeatable spaces.'),
]

for title_text, desc in appearances:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Deck Structure', level=2)

doc.add_paragraph(
    'Each entry in the zone deck links a specific bubble (room) to a specific zone and '
    'assigns its Appearances value. A single room can appear in multiple zone deck entries '
    'if it is eligible for multiple zones. For example, a generic mining tunnel might have '
    'entries in the deck for three different mining-themed zones, each marked as Multiple, '
    'meaning copies of that tunnel can appear in any of those zones.'
)

doc.add_paragraph(
    'The zone deck also respects filters defined in DT_Moria_ZoneBubbleFilters and '
    'DT_Moria_ZoneFilters. These filter files add additional constraints on which rooms '
    'can appear in which zones, acting as whitelists or blacklists on top of the basic '
    'deck assignments.'
)

doc.add_heading('Changing Room Frequency', level=2)

doc.add_paragraph(
    'To change how often a specific room appears in the game, you modify its zone deck '
    'entries. Here are common scenarios:'
)

scenarios = [
    'To guarantee a room always appears: Change its Appearances enum to "Required" in '
    'the zone deck entry for the desired zone.',
    'To prevent a room from appearing: Remove its zone deck entry entirely, or change '
    'its Appearances to "Single" and reduce the zone\'s TargetBubbles so it is less '
    'likely to be drawn.',
    'To increase a room\'s frequency: Ensure it has a "Multiple" Appearances value and '
    'add it to additional zone deck entries in other compatible zones.',
    'To limit a room to one occurrence: Change its Appearances from "Multiple" to '
    '"Single" in all zone deck entries.',
    'To add a room to a new zone: Create a new zone deck entry linking the room to the '
    'target zone with the desired Appearances value.',
]

for i, scenario in enumerate(scenarios, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(scenario)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — Layout Connections
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Layout Connections: The Map Topology', level=1)

doc.add_paragraph(
    'After rooms have been selected, the generator needs to physically connect them into '
    'a navigable map. This is controlled by DT_Moria_LayoutConnections, which defines a '
    'graph of possible connections between bubble types. Think of it as a subway map: each '
    'station (room) has lines (connections) to other stations, and the generator builds '
    'the actual map by following these potential connections.'
)

doc.add_paragraph(
    'Each connection entry specifies:'
)

items = [
    ('Source Bubble', 'The room where the connection starts.'),
    ('Target Bubble', 'The room where the connection leads.'),
    ('Connection Type', 'The kind of passage between them: a doorway, a corridor, a '
     'vertical shaft, a bridge, a highway, or other transition types. Different connection '
     'types may have different visual presentations and navigation requirements.'),
    ('Directionality', 'Whether the connection is one-way or bidirectional. Most '
     'connections are bidirectional (you can travel both ways), but some special connections '
     'like one-way drops or locked doors may be directional.'),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'The connection graph is extensive, with over 70 defined edges between different room '
    'types. This large number of possible connections is what gives the procedural '
    'generation its variety: the same set of rooms can be connected in many different '
    'configurations, producing different map layouts each time.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6 — Layout Constraints
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Layout Constraints: Enforcing Playability', level=1)

doc.add_paragraph(
    'While the connection graph defines what is possible, the constraint system defined '
    'in DT_Moria_LayoutConstraints ensures that certain requirements are always met. '
    'Constraints are the safety net that guarantees the generated world is always '
    'completable, no matter how the random generation plays out.'
)

doc.add_paragraph(
    'Constraints enforce things like:'
)

constraints = [
    'There must always be a path from the starting area to the final area (the game must '
    'be winnable).',
    'Certain bottleneck rooms must appear on the critical path (like the Bridge of '
    'Khazad-dum, which is a required story beat).',
    'Alternate routes must exist in some zones to prevent dead-end frustration.',
    'Certain room types must be reachable before others (for example, a forge must be '
    'accessible before areas that require forged items).',
    'Maximum distances between certain types of rooms (so the player never has to '
    'backtrack excessively to reach a needed facility).',
]

for constraint in constraints:
    doc.add_paragraph(constraint, style='List Bullet')

doc.add_paragraph(
    'The constraint system works in tandem with the connection graph. After the generator '
    'builds an initial layout from the connection graph, the constraint system validates '
    'it. If constraints are not met, the generator adjusts the layout by adding or '
    'rearranging connections until all requirements are satisfied.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7 — Biomes
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Biomes: Environmental Theming', level=1)

doc.add_paragraph(
    'Biomes add environmental character to the generated world. Each zone is associated '
    'with a biome that controls the visual and audio atmosphere: the color of the rock '
    'walls, the ambient sound effects, the types of flora and fauna that appear, the '
    'lighting mood, and other environmental details.'
)

doc.add_paragraph(
    'Biomes are defined in DT_Moria_Biomes and include settings for:'
)

biome_settings = [
    ('Visual Style', 'Rock textures, moss growth, crystal formations, water presence, '
     'general color palette.'),
    ('Lighting Mood', 'Ambient light color and intensity, fog settings, how torches and '
     'other light sources interact with the environment.'),
    ('Audio Atmosphere', 'Background ambient sounds, reverb settings, echo characteristics '
     'that change based on how deep or how large the spaces are.'),
    ('Environmental Hazards', 'What dangers are present: flooding, gas, unstable ceilings, '
     'darkness levels.'),
    ('Resource Distribution', 'What materials can be found in this biome: mithril veins '
     'only appear in specific biomes, certain types of stone or ore are biome-specific.'),
]

for title_text, desc in biome_settings:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'From a modding perspective, biomes are interesting because they affect the player\'s '
    'perception of variety. Even if two zones use similar room layouts, different biome '
    'assignments make them feel distinct. Modifying biome assignments in DT_Moria_Zones '
    'can change the visual character of entire regions of the game.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8 — Zone Templates
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Zone Templates: Pre-Built Layouts', level=1)

doc.add_paragraph(
    'While most of the world generation is procedural, the system also supports pre-built '
    'zone templates defined in DT_Moria_ZoneTemplates. These are hand-designed layouts '
    'that the generator can use as-is instead of building a zone from scratch. Templates '
    'provide a way for the game designers to ensure that certain key areas always have a '
    'specific configuration, even within the random generation system.'
)

doc.add_paragraph(
    'A zone template specifies a fixed arrangement of rooms and connections. The generator '
    'can choose to use a template for a zone instead of building it procedurally, '
    'guaranteeing a known-good layout for critical story areas. This is particularly '
    'useful for the opening areas of each chapter, where the designers want a specific '
    'player experience on first entry.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9 — Filters
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Filters: Fine-Tuning Room Eligibility', level=1)

doc.add_paragraph(
    'On top of the zone deck system, two additional filter files provide fine-grained '
    'control over which rooms can appear where:'
)

filter_items = [
    ('Zone Filters (DT_Moria_ZoneFilters)', 'Group zones into categories and apply '
     'rules at the category level. For example, all "mining" zones might share a common '
     'filter that ensures certain mining-specific rooms are always available, while '
     'excluding rooms that only make sense in urban or forge zones.'),
    ('Zone Bubble Filters (DT_Moria_ZoneBubbleFilters)', 'Apply per-bubble rules that '
     'control which specific rooms are allowed or forbidden in which zones. This acts as '
     'a fine-grained whitelist/blacklist layer on top of the zone deck. A room might be '
     'in the deck for a zone but excluded by a bubble filter due to thematic or gameplay '
     'constraints.'),
]

for title_text, desc in filter_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'The filter system provides a powerful tool for modders who want to restrict or expand '
    'where rooms can appear without modifying the zone deck directly. By adjusting filters, '
    'you can make rooms eligible for zones they normally would not appear in, or prevent '
    'rooms from appearing in zones where they do not fit thematically.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 10 — Bubble Definitions (BF_BB_*)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Bubble Definitions: The Room\'s Identity Card', level=1)

doc.add_paragraph(
    'Every room in the game has two files: a BubbleData file (BD_BB_*) that contains '
    'its physical contents (meshes, breakables, decorations), and a Bubble Definition '
    'file (BF_BB_*) that describes its identity and how it participates in world '
    'generation. The Bubble Definition is the room\'s identity card — it tells the '
    'generation system everything it needs to know to decide whether to place this room '
    'and how to connect it to others.'
)

doc.add_paragraph(
    'A Bubble Definition contains:'
)

bf_fields = [
    ('Bubble Name and Type', 'The internal name (like "BB_Chapter2_GameStart") and the '
     'type classification (what category of room it is).'),
    ('Interfaces', 'The connection points where this room can link to other rooms. Each '
     'interface has a type (Closed, Interior, Large, Tall) and a direction (North, South, '
     'East, West, Up, Down). A room with a "Large" interface on its north side can connect '
     'to another room that has a compatible interface on its south side.'),
    ('Subcells', 'The spatial grid that defines the room\'s footprint. Each subcell is a '
     'unit of space with properties like UsageCategory (EmptyReserved, FullyNavigable) '
     'that tell the game how players can move through the space.'),
    ('Container Proxies', 'Placeholder positions for loot containers, resource nodes, and '
     'other interactive objects that the game populates at runtime.'),
    ('Transporter Pads', 'Positions for the fast-travel pads that allow players to move '
     'between distant rooms. Each pad has a type (Standard) that determines its behavior.'),
    ('Landmark Association', 'A reference to the landmark system (like '
     '"World.Landmark.Chapter2.GameStart") that ties the room to the game\'s waypoint '
     'and navigation systems.'),
    ('Visual Map Style', 'How the room appears on the player\'s in-game map (Cavernous, '
     'Urban, etc.).'),
    ('Difficulty Settings', 'Difficulty overrides that scale enemy strength and resource '
     'availability for this specific room.'),
    ('Biome Association', 'Which biome\'s visual and audio theme applies to this room.'),
    ('Probability Weight', 'A float value on proxy entries that can influence how likely '
     'this room is to be selected when multiple eligible rooms are available.'),
]

for title_text, desc in bf_fields:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 11 — The Game World Catalog
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('The Game World Catalog: Master Registry', level=1)

doc.add_paragraph(
    'BC_GameWorldCatalog is the master registry that ties everything together. It is a '
    'central index of every Bubble Definition in the game. When the world generator needs '
    'to know what rooms are available, it consults this catalog. Every room that can '
    'appear in the game must be registered here.'
)

doc.add_paragraph(
    'The catalog contains references to all BubbleDef files (BF_BB_*) as indexed entries. '
    'Each entry is an ObjectProperty that points to a specific BubbleDef asset. The '
    'generator iterates through the catalog to build its pool of available rooms, then '
    'applies the zone deck, filters, and constraints to determine which rooms actually '
    'get placed.'
)

doc.add_paragraph(
    'From a modding perspective, the Game World Catalog is the gatekeeping file. If you '
    'create a new room, it will not appear in the game until you add it to this catalog. '
    'Conversely, removing a room from this catalog effectively removes it from world '
    'generation entirely (though the BubbleData and BubbleDef files would still exist on '
    'disk).'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 12 — Creating a New Level
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Creating a New Level: The Complete Checklist', level=1)

doc.add_paragraph(
    'Adding a completely new room to Return to Moria\'s world generation requires '
    'creating or modifying several interconnected files. Here is the full process:'
)

doc.add_heading('Step 1: Create the BubbleData (BD_BB_*)', level=2)

doc.add_paragraph(
    'The BubbleData file is the physical content of your room. It needs to contain the '
    'three catalogs (Instanced Mesh Catalog, Instanced Breakable Catalog, and '
    'Construction Catalog) populated with the meshes and objects that make up your room\'s '
    'visual appearance. This is the most labor-intensive step because it defines every '
    'single object in the room and its exact position, rotation, and scale.'
)

doc.add_paragraph(
    'The BubbleData file must follow the game\'s binary format (.uasset) and reference '
    'existing game assets (meshes, materials, collision profiles) through its Import '
    'Table. The file name must follow the BD_BB_ naming convention.'
)

doc.add_heading('Step 2: Create the Bubble Definition (BF_BB_*)', level=2)

doc.add_paragraph(
    'The Bubble Definition tells the generation system about your room. You need to '
    'define:'
)

bf_steps = [
    'The room\'s internal name and type classification.',
    'The interface points (connection doorways) and their directions, so the generator '
    'knows how to attach your room to adjacent rooms.',
    'The subcell grid defining the room\'s spatial footprint, so the game knows what '
    'space your room occupies and where players can walk.',
    'Container proxy positions for loot and resource spawns.',
    'Transporter pad positions if your room should have fast-travel.',
    'The visual map style for the in-game map display.',
    'Difficulty settings appropriate for the intended game phase.',
    'Biome association for environmental theming.',
]

for step in bf_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Step 3: Register in the Game World Catalog', level=2)

doc.add_paragraph(
    'Your new BubbleDef must be added to BC_GameWorldCatalog so the generator knows it '
    'exists. This means adding a new ObjectProperty entry to the catalog\'s export array '
    'that points to your BubbleDef file. Without this registration, the generator will '
    'never consider your room for placement.'
)

doc.add_heading('Step 4: Add to the Zone Deck', level=2)

doc.add_paragraph(
    'Create one or more entries in DT_Moria_ZoneDeck that assign your room to specific '
    'zones. For each zone where your room should be eligible to appear, add a deck entry '
    'with the appropriate Appearances value:'
)

deck_options = [
    'Use "Required" if the room must always appear (use sparingly — too many required '
    'rooms can overcrowd a zone).',
    'Use "Single" for a unique room that should appear at most once per playthrough.',
    'Use "Multiple" for a generic room that adds variety and can appear several times.',
]

for option in deck_options:
    doc.add_paragraph(option, style='List Bullet')

doc.add_heading('Step 5: Update Layout Connections (Optional)', level=2)

doc.add_paragraph(
    'If your room uses interface types that are already well-connected in the layout '
    'graph, the generator may be able to connect it automatically. However, for best '
    'results, add explicit connection entries in DT_Moria_LayoutConnections that define '
    'how your room connects to existing rooms. This ensures the generator can always '
    'find a valid placement for your room.'
)

doc.add_heading('Step 6: Update Filters (If Needed)', level=2)

doc.add_paragraph(
    'If any zone filters or bubble filters would exclude your room, update them to allow '
    'your room in the desired zones. Check both DT_Moria_ZoneFilters and '
    'DT_Moria_ZoneBubbleFilters to ensure compatibility.'
)

doc.add_heading('Step 7: Package and Test', level=2)

doc.add_paragraph(
    'Package all modified files (the new BD_BB_*, new BF_BB_*, modified '
    'BC_GameWorldCatalog, modified DT_Moria_ZoneDeck, and any other changed DataTables) '
    'using the retoc to-zen pipeline to create an IoStore mod pak. Load the game and '
    'verify that your room appears in the generated world, connects properly to adjacent '
    'rooms, and does not cause crashes or navigation issues.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 13 — Probability and Frequency Tuning
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Probability and Frequency Tuning: A Practical Guide', level=1)

doc.add_paragraph(
    'One of the most accessible forms of world generation modding is adjusting how '
    'frequently certain rooms, zones, or biomes appear. Here is a summary of every lever '
    'available and what it controls:'
)

# Table of controls
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'File'
hdr_cells[1].text = 'Field'
hdr_cells[2].text = 'Type'
hdr_cells[3].text = 'What It Controls'

controls = [
    ('DT_Moria_ZoneDeck', 'Appearances', 'Enum', 'Required = guaranteed; Single = max once; Multiple = repeatable'),
    ('DT_Moria_Zones', 'TargetBubbles', 'Integer', 'Target number of rooms in a zone'),
    ('DT_Moria_Zones', 'NewBubbleChance', 'Float', 'Probability of adding another room (0.0-1.0)'),
    ('DT_Moria_Zones', 'AdditionalOpeningChance', 'Float', 'Probability of extra connections between rooms'),
    ('DT_Moria_Zones', 'GenerationPriority', 'Integer', 'Build order — higher = built first'),
    ('BF_BB_* BubbleDefs', 'Probability', 'Float', 'Weight on proxy entries for selection likelihood'),
]

for file, field, dtype, desc in controls:
    row_cells = table.add_row().cells
    row_cells[0].text = file
    row_cells[1].text = field
    row_cells[2].text = dtype
    row_cells[3].text = desc

doc.add_paragraph()

doc.add_paragraph(
    'These controls interact with each other. For example, setting a room to "Multiple" '
    'in a zone with a high TargetBubbles and high NewBubbleChance will make that room '
    'appear frequently. Conversely, setting it to "Single" in a zone with low '
    'TargetBubbles means it will rarely appear.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 14 — Key Files Reference
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Key Files Reference', level=1)

doc.add_paragraph(
    'All world generation configuration files are located in Tech/Data/GameWorld/ within '
    'the game\'s data. Here is a complete reference of every file involved:'
)

files = [
    ('BP_Moria_WorldLayout.uasset', 'The main world layout Blueprint — the code that '
     'reads all the DataTables and actually runs the generation algorithm.'),
    ('DT_Moria_Chapters.uasset', 'Chapter progression definitions. Controls story flow.'),
    ('DT_Moria_Zones.uasset', 'Zone definitions with size, probability, and priority '
     'parameters.'),
    ('DT_Moria_ZoneDeck.uasset', 'The room selection deck. Controls which rooms can '
     'appear in which zones and with what frequency.'),
    ('DT_Moria_ZoneTemplates.uasset', 'Pre-built zone layouts for hand-designed areas.'),
    ('DT_Moria_LayoutConnections.uasset', 'The connection graph defining how rooms can '
     'link to each other.'),
    ('DT_Moria_LayoutConstraints.uasset', 'Constraints that enforce critical path '
     'requirements and playability.'),
    ('DT_Moria_ZoneFilters.uasset', 'Zone-level filtering rules for room eligibility.'),
    ('DT_Moria_ZoneBubbleFilters.uasset', 'Per-bubble filtering rules.'),
    ('DT_Moria_Biomes.uasset', 'Biome definitions for visual and audio theming.'),
    ('DT_Moria_Landmarks.uasset', 'Landmark definitions for navigation and waypoints.'),
    ('BC_GameWorldCatalog.uasset', 'Master registry of all BubbleDef files.'),
    ('BF_BB_*.uasset files', 'Individual Bubble Definition files in '
     'Tech/Data/BubbleDefs/GameWorldCatalog/.'),
    ('BD_BB_*.uasset files', 'Individual BubbleData files in '
     'Tech/Data/Bubbles/GameWorldCatalog/.'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'File'
hdr_cells[1].text = 'Purpose'
for file, purpose in files:
    row_cells = table.add_row().cells
    row_cells[0].text = file
    row_cells[1].text = purpose

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 15 — Data Flow Summary
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Data Flow: How Generation Proceeds', level=1)

doc.add_paragraph(
    'To tie everything together, here is the step-by-step data flow when the game '
    'generates a new world:'
)

flow_steps = [
    'The game reads DT_Moria_Chapters to determine the chapter progression and which '
    'zones belong to each chapter.',
    'For each chapter, it reads the zone definitions from DT_Moria_Zones to get the '
    'generation parameters (TargetBubbles, NewBubbleChance, etc.).',
    'For each zone, it consults DT_Moria_ZoneDeck to build the deck of eligible rooms.',
    'It applies filters from DT_Moria_ZoneFilters and DT_Moria_ZoneBubbleFilters to '
    'trim the deck down to truly eligible rooms.',
    'It checks DT_Moria_ZoneTemplates to see if a pre-built layout should be used '
    'instead of procedural generation.',
    'For procedurally generated zones, it draws rooms from the deck according to the '
    'Appearances rules (Required rooms are placed first, then Single and Multiple rooms '
    'fill the remaining target count).',
    'It uses DT_Moria_LayoutConnections to determine how the selected rooms connect to '
    'each other, building the physical topology of the zone.',
    'It validates the layout against DT_Moria_LayoutConstraints to ensure all critical '
    'paths exist and the game is completable.',
    'It applies biome settings from DT_Moria_Biomes to each zone for environmental '
    'theming.',
    'Finally, for each placed room, it loads the corresponding BubbleData file (BD_BB_*) '
    'and instantiates all the objects from its catalogs into the game world.',
]

for i, step in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

doc.add_paragraph()

doc.add_paragraph(
    'This entire process happens once when a new world is created. The generated layout '
    'is then saved and remains fixed for that playthrough. The procedural generation '
    'provides variety between playthroughs, while the constraint system ensures that '
    'every generated world is completable and hits all the required story beats.'
)

# ── Save ──
output_path = Path(__file__).parent / 'How World Generation Works in Return to Moria.docx'
doc.save(str(output_path))
print(f'Saved: {output_path}')
print(f'Sections: 15')
print(f'Paragraphs: {len(doc.paragraphs)}')
