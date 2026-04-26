"""
Generate the MASTER Moria WorldGen document — single source of truth.

Combines:
  1. Title & data-model overview (originally docs/worldgen-data-model.docx)
  2. Architecture diagram + explanatory sections (generate_worldgen_architecture_doc.py)
  3. Full inventory of every DataTable row (generate_worldgen_inventory.py)

Strategy:
  - Re-run the two existing generator scripts so their standalone outputs stay
    current as an artifact.
  - Build a small cover/overview document with python-docx.
  - Use docxcompose to stitch:   [overview]  +  [data-model]  +  [architecture]  +  [inventory]
    into docs/Moria_WorldGen_Master.docx.

Run:
    python scripts/generate_worldgen_master_doc.py
"""

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor
from docxcompose.composer import Composer


SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DOCS = PROJECT_ROOT / 'docs'
DOCS.mkdir(exist_ok=True)

DATA_MODEL_DOC = DOCS / 'worldgen-data-model.docx'
ARCH_DOC = DOCS / 'Moria_WorldGen_Architecture.docx'
INV_DOC = DOCS / 'Moria_WorldGen_Inventory.docx'
LEVELS_DOC = DOCS / 'moria-levels-complete.docx'
MASTER_DOC = DOCS / 'Moria_WorldGen_Master.docx'

GENERATORS = [
    SCRIPT_DIR / 'generate_worldgen_architecture_doc.py',
    SCRIPT_DIR / 'generate_worldgen_inventory.py',
]


# -----------------------------------------------------------------------------
# Step 1 — refresh the component documents so the master reflects the latest
# code + data state.
# -----------------------------------------------------------------------------

def refresh_component_docs():
    for gen in GENERATORS:
        if not gen.exists():
            print(f'WARN: missing generator {gen}')
            continue
        print(f'[refresh] running {gen.name}')
        r = subprocess.run([sys.executable, str(gen)],
                            capture_output=True, text=True)
        if r.returncode != 0:
            print(f'  stderr: {r.stderr[:400]}')
            raise SystemExit(f'{gen.name} failed')


# -----------------------------------------------------------------------------
# Step 2 — build the master cover / overview sheet that sits in front of the
# existing content. Keeps the master document self-describing so a reader
# doesn't need to know the generator script existed.
# -----------------------------------------------------------------------------

def build_cover(path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)

    # Cover / data-model / architecture are portrait. The inventory doc
    # brings its own landscape section break, so this portrait section
    # ends when the inventory chunk starts.
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Moria WorldGen — Master Reference')
    r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x2A, 0x44)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run('Data model · Architecture · Inventory')
    sr.italic = True; sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    h = doc.add_heading('What is in this document', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph(
        'This is the single source of truth for the Moria WorldGen editor '
        'project. It combines three complementary views:')

    bullets = [
        ('Part I — WorldGen Data Model',
         'The tier diagram: World & Config → Chapters → Zones → Decks → '
         'Bubbles, plus support tables that constrain layout, navigation, '
         'and points of interest.'),
        ('Part II — Architecture & Editor',
         'The seven DataTables, field-level relationships, editor tabs, '
         'build pipeline, gotchas, bubble classification, special flags, '
         'zone sizing, landmark connections, and common confusions.'),
        ('Part III — Full Inventory',
         'Every chapter, zone, biome, landmark, deck, and filter with '
         'cross-references, generated directly from the current JSON state.'),
    ]
    for title, body in bullets:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(title)
        r.bold = True
        p.add_run(': ' + body)

    doc.add_paragraph()

    h = doc.add_heading('How to regenerate this document', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph(
        'This document is assembled automatically from the editor code and '
        'current DataTable JSONs. To rebuild it after edits:')
    p = doc.add_paragraph()
    r = p.add_run('    python scripts/generate_worldgen_master_doc.py')
    r.font.name = 'Consolas'
    r.font.size = Pt(10)

    doc.add_paragraph(
        'The script re-runs the architecture and inventory generators, then '
        'uses docxcompose to stitch this cover page together with the '
        'worldgen-data-model overview, the architecture reference, and the '
        'inventory. The intermediate standalone docs '
        '(Moria_WorldGen_Architecture.docx and Moria_WorldGen_Inventory.docx) '
        'remain available as separate deliverables.')

    doc.add_page_break()
    doc.save(path)


# -----------------------------------------------------------------------------
# Step 3 — stitch the components together with docxcompose.
# -----------------------------------------------------------------------------

def build_master():
    # Cover lives in a temp file on disk because Composer works with file paths.
    tmp_cover = DOCS / '_tmp_master_cover.docx'
    build_cover(tmp_cover)

    print(f'[compose] cover     = {tmp_cover.name}')
    master = Document(tmp_cover)
    composer = Composer(master)

    to_append = []
    for src in (DATA_MODEL_DOC, ARCH_DOC, INV_DOC):
        if not src.exists():
            print(f'WARN: {src.name} not found — skipping')
            continue
        to_append.append(src)

    for src in to_append:
        print(f'[compose] append    = {src.name}')
        composer.append(Document(src))

    composer.save(MASTER_DOC)
    try:
        tmp_cover.unlink()
    except OSError:
        pass

    print(f'\nWrote {MASTER_DOC}')
    sz = MASTER_DOC.stat().st_size
    print(f'  size: {sz:,} bytes')


def main():
    refresh_component_docs()
    build_master()


if __name__ == '__main__':
    main()
