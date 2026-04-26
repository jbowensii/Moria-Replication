"""
Generate a Word document and PNG diagram for Moria WorldGen Architecture.

Outputs:
    docs/worldgen_architecture_diagram.png
    docs/Moria_WorldGen_Architecture.docx
"""

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.lines import Line2D

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DOCS_DIR = PROJECT_ROOT / 'docs'
DOCS_DIR.mkdir(exist_ok=True)
PNG_PATH = DOCS_DIR / 'worldgen_architecture_diagram.png'
DOCX_PATH = DOCS_DIR / 'Moria_WorldGen_Architecture.docx'


# -----------------------------------------------------------------------------
# 1. Render the diagram
# -----------------------------------------------------------------------------

COLORS = {
    'zones':     '#4a90e2',
    'chapters':  '#7ed321',
    'biomes':    '#f5a623',
    'decks':     '#bd10e0',
    'filters':   '#50e3c2',
    'landmarks': '#d0021b',
    'bubbles':   '#9b9b9b',
    'bp':        '#c7ced4',
}


def draw_box(ax, x, y, w, h, title, subtitle, color):
    """Draw a rounded box with a coloured header bar."""
    # Main body (white)
    body = FancyBboxPatch(
        (x, y), w, h, boxstyle='round,pad=0.02,rounding_size=0.12',
        linewidth=1.5, edgecolor='#333', facecolor='white', zorder=2)
    ax.add_patch(body)
    # Header bar
    header_h = 0.38
    header = FancyBboxPatch(
        (x, y + h - header_h), w, header_h,
        boxstyle='round,pad=0.02,rounding_size=0.12',
        linewidth=0, facecolor=color, zorder=3)
    ax.add_patch(header)
    # Mask rectangle for bottom of header (so only the top is rounded)
    ax.add_patch(plt.Rectangle(
        (x, y + h - header_h), w, 0.12,
        linewidth=0, facecolor=color, zorder=4))

    ax.text(x + w/2, y + h - header_h/2, title,
            ha='center', va='center', fontsize=10, fontweight='bold',
            color='white', zorder=5)
    ax.text(x + w/2, y + h - header_h - 0.18, subtitle,
            ha='center', va='top', fontsize=8.5, color='#333', zorder=5,
            wrap=True)


def arrow(ax, x1, y1, x2, y2, label='', curve=0.0, color='#555', offset=(0, 0)):
    """Draw a labelled arrow between two points."""
    style = f'arc3,rad={curve}'
    a = FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle='-|>', mutation_scale=14, linewidth=1.4,
        color=color, connectionstyle=style, zorder=6)
    ax.add_patch(a)
    if label:
        mx = (x1 + x2) / 2 + offset[0]
        my = (y1 + y2) / 2 + offset[1]
        ax.text(mx, my, label, fontsize=7.5, color=color,
                bbox=dict(boxstyle='round,pad=0.15', fc='white',
                          ec='none', alpha=0.92),
                ha='center', va='center', zorder=7)


def render_diagram():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('#f5f6f8')

    # Title
    ax.text(7, 8.65, 'Moria WorldGen Architecture',
            ha='center', fontsize=17, fontweight='bold', color='#222')
    ax.text(7, 8.25,
            'Seven DataTables that build the sandbox world — and how they link',
            ha='center', fontsize=10, color='#555', style='italic')

    # Centre: Zones (the hub)
    draw_box(ax, 5.5, 4.5, 3.0, 1.5, 'DT_Moria_Zones',
             '91 rows  •  44 SandboxSmall\nevery cell of the world',
             COLORS['zones'])

    # Top: Chapters
    draw_box(ax, 5.5, 7.1, 3.0, 1.05, 'DT_Moria_Chapters',
             '24 rows  •  Z-layers + enemy scaling',
             COLORS['chapters'])

    # Left: Biomes
    draw_box(ax, 0.6, 5.0, 3.0, 1.3, 'DT_Moria_Biomes',
             '34 rows  •  audio / deco /\nrock / atmospherics presets',
             COLORS['biomes'])

    # Right: ZoneDeck
    draw_box(ax, 10.4, 5.0, 3.0, 1.3, 'DT_Moria_ZoneDeck',
             '67 rows  •  pools of bubbles\n(Required / Single / Multiple)',
             COLORS['decks'])

    # Bottom: Landmarks
    draw_box(ax, 5.5, 2.1, 3.0, 1.3, 'DT_Moria_Landmarks',
             '70 rows  •  anchors +\nGuaranteedConnections graph',
             COLORS['landmarks'])

    # Right of landmarks: Bubble names
    draw_box(ax, 10.4, 2.3, 3.0, 1.0, 'BB_* bubbles',
             'Physical room uassets\n(BubbleData + BubbleDef)',
             COLORS['bubbles'])

    # Left of landmarks: ZoneBubbleFilters
    draw_box(ax, 0.6, 2.3, 3.0, 1.0, 'DT_Moria_ZoneBubbleFilters',
             '12 rows  •  whitelist / blacklist\nfor bubbles',
             COLORS['filters'])

    # Bottom-left: Blueprint Data Assets (targets of Biome refs)
    draw_box(ax, 0.6, 0.3, 3.0, 1.0, 'Blueprint DataAssets',
             'DA_Amb_*, DA_Deco_*,\nDA_Rock_*, DA_Atmos_*',
             COLORS['bp'])

    # Arrows (all labelled by field)
    # Zones -> Chapters
    arrow(ax, 7.0, 6.0, 7.0, 7.1, 'Chapter (RowHandle)', curve=0)

    # Zones -> Biomes
    arrow(ax, 5.5, 5.0, 3.6, 5.5, 'Biome (GameplayTag)', curve=-0.15)

    # Biomes -> BPs
    arrow(ax, 2.1, 5.0, 2.1, 1.3, 'AudioConfig,\nDecoConfig, etc.',
          curve=0.15, offset=(0.1, 0))

    # Zones -> ZoneDeck
    arrow(ax, 8.5, 5.5, 10.4, 5.5, 'BubbleDeck /\nPassageDeck', curve=0.15)

    # ZoneDeck -> Bubbles
    arrow(ax, 11.9, 5.0, 11.9, 3.3, 'DeckEntries[].Bubble', curve=0)

    # Zones -> Landmarks
    arrow(ax, 7.0, 4.5, 7.0, 3.4, 'LandmarkHandles[]', curve=0)

    # Landmarks -> Landmarks (self-loop for GuaranteedConnections)
    ax.add_patch(FancyArrowPatch(
        (8.3, 2.9), (8.65, 2.5),
        connectionstyle='arc3,rad=1.2',
        arrowstyle='-|>', mutation_scale=14,
        color=COLORS['landmarks'], linewidth=1.4, zorder=6))
    ax.text(9.0, 2.6, 'GuaranteedConnections\n(landmark ↔ landmark)',
            fontsize=7.5, color=COLORS['landmarks'],
            bbox=dict(boxstyle='round,pad=0.2', fc='white', ec='none', alpha=0.92))

    # Filters -> Bubbles (applies to)
    arrow(ax, 3.6, 2.8, 10.4, 2.8, 'applies whitelist / blacklist',
          curve=0.15, color='#268080', offset=(0, 0.15))

    # Legend / key
    ax.text(0.2, 0.0, '→ Solid arrow = RowHandle / tag reference',
            fontsize=7.5, color='#555')
    ax.text(0.2, -0.22, '…  Filter arrow = applied during bubble selection',
            fontsize=7.5, color='#555')

    plt.tight_layout()
    fig.savefig(PNG_PATH, dpi=160, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    print(f'Wrote {PNG_PATH}')


# -----------------------------------------------------------------------------
# 2. Generate the Word document
# -----------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
        set_cell_shading(hdr[i], '2E5597')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p


def build_docx():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)

    # Trim margins slightly so 6.8"-wide tables fit in portrait US Letter.
    # Also keeps the embedded diagram centered without clipping.
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Cover title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('Moria WorldGen Architecture')
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x22, 0x2A, 0x44)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Reference for the Moria WorldGen Editor')
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    # Diagram
    if PNG_PATH.exists():
        doc.add_picture(str(PNG_PATH), width=Inches(6.5))
        cap = doc.add_paragraph('Figure 1 — The seven DataTables and how they link.')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Section — what each DataTable does
    add_heading(doc, 'The seven DataTables — what each one does', level=1)
    add_table(doc,
        headers=['DataTable', 'Rows', 'Role', 'Editor tab'],
        rows=[
            ['DT_Moria_Zones', '91',
             'Cell grid of the playable world. Each zone holds a biome, chapter, landmarks, deck references, and generation tuning.',
             'Zones'],
            ['DT_Moria_Chapters', '24',
             'Z-layer grouping + enemy scaling bucket. How the game slices the world vertically.',
             'Chapters'],
            ['DT_Moria_Biomes', '34',
             'Audio / lighting / deco / rock / atmospherics presets. The theme of a region.',
             'Biomes'],
            ['DT_Moria_ZoneDeck', '67',
             'Pools of bubbles with Appearances rules (Required / Single / Multiple) and zone-entrance flags.',
             'Bubbles'],
            ['DT_Moria_ZoneBubbleFilters', '12',
             'Extra allow / deny lists layered on top of deck selection.',
             'Filters'],
            ['DT_Moria_Landmarks', '70',
             'Named anchor points. Each row bundles a bubble name, position, connections list, and display data.',
             'Landmarks'],
            ['(synthesised)', '–',
             'Mappings tab joins all of the above into one row-per-zone table.',
             'Mappings'],
            ['(drawing)', '–',
             'Map tab renders zones as isometric boxes, one chapter at a time, with connectivity overlay.',
             'Map'],
        ],
        col_widths=[2.0, 0.6, 3.4, 0.8])

    # Section — Zone field details
    add_heading(doc, 'Zone fields and what they link to', level=1)
    add_table(doc,
        headers=['Zone field', 'Type', 'Links to'],
        rows=[
            ['ZoneSet', 'enum', 'SandboxSmall | SandboxMedium | Moria | Expedition'],
            ['Chapter', 'RowHandle', 'DT_Moria_Chapters row'],
            ['Biome', 'GameplayTag', 'World.Biome.* (conceptually DT_Moria_Biomes row)'],
            ['BubbleDeck', 'RowHandle', 'DT_Moria_ZoneDeck row (bubbles the zone spawns)'],
            ['PassageDeck', 'RowHandle', 'DT_Moria_ZoneDeck row (corridor bubbles)'],
            ['Challenges', 'RowHandle', 'DT_Moria_ZoneChallenges row'],
            ['StagingContainers[]', 'Array', 'DT_Moria_ResourceContainers rows (barrels, bones…)'],
            ['Templates[]', 'Array', 'DT_Moria_ZoneTemplates rows (when bUseTemplate=true)'],
            ['LandmarkHandles[]', 'Array',
             'DT_Moria_Landmarks rows + Placement + bExtendedConnectivityLandmark'],
            ['Position', 'IntVector(X,Y,Z)', 'Zone origin in bubble-grid units'],
            ['TargetSize', 'IntVector(X,Y,Z)', 'Zone footprint in bubble-grid units'],
            ['VisualMapStyle', 'enum', 'Urban | Cavernous | Secret | Outside'],
            ['ZoneTemperature / Water / LightPrevalence', 'floats', 'Environmental tuning'],
            ['TargetBubbles', 'int', 'How many bubbles to try to place'],
            ['NewBubbleChance', 'float', 'Probability of additional bubble spawns'],
        ],
        col_widths=[1.7, 1.2, 3.9])

    # Section — Chapter
    add_heading(doc, 'Chapter fields', level=1)
    add_table(doc,
        headers=['Field', 'Meaning'],
        rows=[
            ['Name (row key)', 'Moria-chapter-2, SandboxSmall-chapter-3, SandboxSmall-chapter-9…'],
            ['ZoneSet (enum)', 'Which zone set this chapter belongs to'],
            ['ChapterID (int)', 'Numeric chapter ID'],
            ['DisplayName', 'String-table key (e.g. Chapter2.Name)'],
            ['Layer (int)', 'Vertical layer ordering'],
            ['EnemyScalingLevel', 'Difficulty tier'],
            ['MinZ / MaxZ / PrimeZ', 'Z-extent for the chapter'],
            ['EnabledState', 'Live / Disabled'],
        ],
        col_widths=[2.0, 4.8])

    # Section — Biome
    add_heading(doc, 'Biome fields', level=1)
    add_table(doc,
        headers=['Field', 'Links to'],
        rows=[
            ['BiomeId', 'Gameplay-tag bridge to World.Biome.*'],
            ['DisplayName', 'String-table lookup'],
            ['AudioConfig', 'DA_Amb_* DataAsset — ambient sound preset'],
            ['DecoConfig', 'DA_Deco_* DataAsset — surface deco rules'],
            ['InteriorDecoConfig', 'DA_Deco_Interior_* — indoor deco'],
            ['GuidemeshMaterial', 'Material — navigation-mesh appearance'],
            ['RockConfig', 'DA_Rock_* — rock meshes + materials'],
            ['AtmosphericsConfig', 'DA_Atmos_* — fog / sky / weather'],
            ['DefaultCellData', 'Map property — per-cell defaults'],
        ],
        col_widths=[2.0, 4.8])

    # Section — ZoneDeck
    add_heading(doc, 'ZoneDeck — the bubble pools', level=1)
    add_table(doc,
        headers=['Entry field', 'Meaning'],
        rows=[
            ['Name (row key)',
             'Sandbox_WestgateBubbles, Sandbox_MinesPassages, …'],
            ['DeckEntries[].Bubble',
             'BB_* bubble uasset name'],
            ['DeckEntries[].Appearances',
             'Required | Single | Multiple'],
            ['DeckEntries[].bZoneEntrance',
             'True if this bubble is the entry point'],
            ['EnabledState', 'Live / Disabled'],
        ],
        col_widths=[2.5, 4.3])

    # Section — Filters
    add_heading(doc, 'ZoneBubbleFilters', level=1)
    add_table(doc,
        headers=['Field', 'Meaning'],
        rows=[
            ['Name (row key)',
             'Default, TestBlacklist, TestWhitelist, …'],
            ['Whitelist[]', 'Bubble names allowed (empty = no restriction)'],
            ['Blacklist[]', 'Bubble names forbidden'],
            ['EnabledState', 'Live / Disabled'],
        ],
        col_widths=[2.0, 4.8])

    # Section — Landmark
    add_heading(doc, 'Landmark fields', level=1)
    add_table(doc,
        headers=['Field', 'Meaning / links'],
        rows=[
            ['Name', 'Chapter2.DoorsOfDurin, Sandbox.Aftermath, …'],
            ['BaseBubbleName', 'BB_* — the physical room uasset for this landmark'],
            ['Placement (enum)', 'Fixed | Free | Sidequest | Unspecified'],
            ['BasePosition', 'World anchor vector (per-chapter coordinate space)'],
            ['GuaranteedConnections[]',
             'GameplayTag "World.Landmark.<row name>" for every peer this landmark must connect to'],
            ['DisplayName / LoreText', 'String-table lookups'],
            ['LandmarkIcon / LandmarkMapIcon', 'UI textures'],
            ['bPlayerStartLocation', 'Spawn here when a run starts'],
            ['ChallengeRating', 'Difficulty hint'],
            ['BiomeOverride', 'Optional biome override for this landmark'],
            ['EnabledState', 'Live / Disabled'],
        ],
        col_widths=[2.4, 4.4])

    # Section — Q&A
    add_heading(doc, 'How the data answers questions', level=1)

    add_heading(doc,
        'Q1. What bubbles can the Westgate zone spawn?', level=2)
    doc.add_paragraph(
        '1. Open zone Sandbox_Small_Suburban_A_Westgate in DT_Moria_Zones.\n'
        '2. Read BubbleDeck → Sandbox_WestgateBubbles.\n'
        '3. Open that row in DT_Moria_ZoneDeck and read DeckEntries[].\n'
        '4. Each entry is a BB_* bubble name with an Appearances rule.\n'
        '5. Apply any DT_Moria_ZoneBubbleFilters whitelist / blacklist on top.\n'
        '→ In the editor: Zones tab → select zone → the "Bubbles that this zone '
        'can spawn" listbox does this walk automatically.')

    add_heading(doc,
        'Q2. Why is the Elven Quarter connected to the Doors of Durin?', level=2)
    doc.add_paragraph(
        '1. Open zone Sandbox_Small_ElvenQuarter_A in DT_Moria_Zones.\n'
        '2. Read LandmarkHandles[] — e.g. Sandbox.ElvenQuarterPromenade.\n'
        '3. Open that landmark in DT_Moria_Landmarks.\n'
        '4. Read GuaranteedConnections[] — each entry "World.Landmark.<row>" points to a peer.\n'
        '5. Peers are landmarks held by other zones; that is the connectivity graph.\n'
        '→ Editor: Landmarks tab shows the peer list, Map tab draws it as green '
        'dashed lines.')

    add_heading(doc,
        'Q3. How does a zone get its mood (lighting / ambient)?', level=2)
    doc.add_paragraph(
        '1. Zone has a Biome GameplayTag (e.g. World.Biome.WesternTown).\n'
        '2. That biome row in DT_Moria_Biomes bundles AudioConfig + DecoConfig + '
        'AtmosphericsConfig + RockConfig + LightingInfo.\n'
        '3. Zone can override with its own ZoneLightingBehavior, LightingInfo, '
        'ZoneTemperature, WaterPrevalence, LightPrevalence.\n'
        '4. A landmark can further override the zone\'s biome via BiomeOverride.')

    add_heading(doc,
        'Q4. How do I add book-accurate Levels 5–7 and the Foundations of Stone?',
        level=2)
    doc.add_paragraph(
        '1. Chapters tab → Add Book-Accurate Chapters (9–15) button — creates '
        'the new rows with default Z values.\n'
        '2. Zones tab → reassign existing zones to the new chapters via the '
        'Chapter dropdown, or duplicate zones and point the copies at the new '
        'chapters.\n'
        '3. Build → Build Mod Pak bundles DT_Moria_Chapters + DT_Moria_Zones '
        'into SandboxMod_P.\n'
        '4. Landmarks / decks / biomes only need editing if you want the new '
        'chapters to look different.')

    # Section — tab map
    add_heading(doc, 'Editor tabs and the DataTables they edit', level=1)
    add_table(doc,
        headers=['Tab', 'Edits', 'Notes'],
        rows=[
            ['Zones', 'DT_Moria_Zones',
             'Bubble Sources section reads DT_Moria_ZoneDeck and shows the '
             'actual bubble list; Landmarks section picks from DT_Moria_Landmarks.'],
            ['Chapters', 'DT_Moria_Chapters',
             'Add Chapter + Add Book-Accurate Chapters (9–15) buttons.'],
            ['Biomes', 'DT_Moria_Biomes',
             'DisplayName + Enabled editable; Object refs resolved from Imports table.'],
            ['Bubbles', 'DT_Moria_ZoneDeck',
             'Add / edit / remove DeckEntries; bubble picker shows known BB_* names.'],
            ['Filters', 'DT_Moria_ZoneBubbleFilters',
             'Whitelist / Blacklist editing per filter row.'],
            ['Landmarks', 'DT_Moria_Landmarks',
             'Edit BaseBubbleName, Placement, ChallengeRating, PlayerStart flag, '
             'and the GuaranteedConnections peer list.'],
            ['Mappings', '(read-only view)',
             'Joins Zone → Chapter / Biome / BubbleDeck / PassageDeck / Landmarks. '
             'Double-click any cell to jump.'],
            ['Map', '(read-only view)',
             'Isometric canvas. One chapter at a time, grid layout so every box is '
             'visible. Green dashed lines show landmark connections.'],
        ],
        col_widths=[1.1, 2.0, 3.7])

    # Section — build pipeline
    add_heading(doc, 'Build pipeline', level=1)
    doc.add_paragraph(
        'The Build Mod Pak button serialises every loaded DataTable JSON back '
        'into a .uasset with UAssetGUI and packages them into a single IoStore '
        'mod pak with retoc.')
    add_mono(doc,
        'JSON (edited)\n'
        '   ↓ UAssetGUI fromjson (VER_UE4_27)\n'
        'uasset + uexp, staged under\n'
        '   Moria/Content/Tech/Data/GameWorld/\n'
        '   ↓ retoc to-zen (UE4_27)\n'
        'SandboxMod_P.pak + .ucas + .utoc\n'
        '   ↓ zip\n'
        '~/Downloads/SandboxMod_v{VERSION}.zip\n\n'
        'Install: extract into the game\'s Paks/~mods/ folder.\n'
        'Uninstall: delete the pak triplet — the game reverts to stock behaviour.')

    # Section — out of scope
    add_heading(doc, 'Out of scope for the current mod pak', level=1)
    doc.add_paragraph(
        'These DataTables and assets are NOT touched by the editor. They could '
        'be added as future tabs if the need arises.')
    add_table(doc,
        headers=['Asset', 'What it is'],
        rows=[
            ['DT_Moria_ZoneTemplates', 'Zone blueprints / shared presets'],
            ['DT_Moria_Waypoints', 'Navigation waypoints'],
            ['DT_Moria_LayoutConnections', 'Bubble-level connection graph'],
            ['DT_Moria_LayoutConstraints', 'Placement constraints on bubbles'],
            ['BC_GameWorldCatalog', 'Master bubble registry'],
            ['BF_BB_* / BD_BB_*', 'Per-bubble definition + instance data'],
            ['DT_Moria_ResourceContainers', 'Barrels, dwarf remains, loot piles'],
            ['DT_Moria_Architecture, DT_Moria_WorldInfo, DT_WorldGen',
             'Misc worldgen config'],
        ],
        col_widths=[2.8, 4.0])

    # --- NEW SECTION: Bubble classification -----------------------------
    add_heading(doc, 'Bubble classification — landmark bubbles vs deck bubbles',
                level=1)
    doc.add_paragraph(
        'Every BB_* bubble in the game plays exactly one role — it is either '
        'a landmark anchor (hand-crafted set-piece) or a deck bubble (generic '
        'reusable room). The two sets never overlap.')
    add_table(doc,
        headers=['Class', 'Count', 'Role', 'Examples', 'Appears in'],
        rows=[
            ['Landmark bubbles', '41',
             'Hand-crafted set-piece rooms tied to named locations',
             'BB_Chapter2_DoorsOfDurin, BB_Sandbox_ElvenForge, BB_Sandbox_Elevator_Urban',
             'DT_Moria_Landmarks.BaseBubbleName only'],
            ['Deck bubbles', '40',
             'Generic rooms shuffled procedurally',
             'BB_DwarfHall1, BB_Highway, BB_Passage_MiningTunnels, BB_UrbanCommunity',
             'DT_Moria_ZoneDeck.DeckEntries[].Bubble only'],
            ['Intersection', '0',
             '(no overlap)', '—', '—'],
        ],
        col_widths=[1.4, 0.5, 1.7, 1.7, 1.4])
    doc.add_paragraph(
        'Important consequence: landmark bubbles spawn ONLY when a landmark '
        'references them. Remove a landmark from every zone and its bubble '
        'becomes completely unreachable in-game.')

    # --- NEW SECTION: Special flags -------------------------------------
    add_heading(doc, 'Special flags', level=1)

    add_heading(doc, 'bExtendedConnectivityLandmark (on zone LandmarkHandles entry)',
                level=2)
    doc.add_paragraph(
        'Marks a landmark as a cross-chapter transit hub. Only 7 entries in '
        'the entire game have this set True — every one is a vertical transit '
        'point:')
    add_table(doc,
        headers=['Zone', 'Landmark', 'Placement'],
        rows=[
            ['Sandbox_Small_Elevator_B', 'Sandbox.FirstStair', 'Fixed'],
            ['Sandbox_Small_Elevator_C', 'Sandbox.SecondStair', 'Fixed'],
            ['Sandbox_Small_Elevator_D', 'Sandbox.ThirdStair', 'Fixed'],
            ['Sandbox_Small_Elevator_E', 'Sandbox.FourthStair', 'Fixed'],
            ['Sandbox_Small_Elevator_F', 'Sandbox.FifthStair', 'Fixed'],
            ['Sandbox_Small_City_A_EasternBastion', 'Sandbox.EasternStairs', 'Fixed'],
            ['Sandbox_Small_Mines_C', 'Sandbox.Deep1MineNexus', 'Interior'],
        ],
        col_widths=[2.8, 2.2, 1.0])
    doc.add_paragraph(
        'Rule: tick it for stairs/elevators/cross-chapter portals. Leave '
        'unchecked for set-pieces, forges, towns, altars, tombs.')

    add_heading(doc, 'bZoneEntrance (on deck entry)', level=2)
    doc.add_paragraph(
        'Marks a bubble as the mandatory gateway where the player arrives '
        'when entering the zone from outside. Only 6 entries across 314 '
        'total deck entries have it set. All on BubbleDecks, none on '
        'PassageDecks.')
    add_table(doc,
        headers=['Deck', 'Entrance bubble', 'Appearances'],
        rows=[
            ['OrcTownBubblesA', 'BB_OrcTown_Gate', 'Required'],
            ['GhashOrcTownBubbles', 'BB_OrcTown_Gate', 'Required'],
            ['DarkestDeepsBubbles', 'BB_DarkestDeepsEntrance', 'Required'],
            ['Sandbox_OrcTownBubbles', 'BB_OrcTown_Gate', 'Required'],
            ['Expedition_DarkestDeepsBubbles', 'BB_DarkestDeepsEntrance', 'Required'],
            ['Expedition_GrendelBubbles', 'BB_TrollCave_AllInterfaces', 'Single'],
        ],
        col_widths=[2.8, 2.6, 1.0])

    # --- NEW SECTION: Zone sizing ---------------------------------------
    add_heading(doc, 'Zone sizing — TargetSize vs TargetBubbles', level=1)
    doc.add_paragraph(
        'TargetSize (X, Y, Z) is the zone\'s maximum bounding box in bubble-'
        'grid units (each unit is roughly one 65m x 65m x 30m room). '
        'TargetBubbles is how many rooms to actually place inside that box. '
        'Volume is always much bigger than the bubble count — the remaining '
        'space is air, stone, or transition.')
    add_table(doc,
        headers=['Zone', 'Size', 'Volume', 'Bubbles', 'Density'],
        rows=[
            ['Westgate',      '(8, 8, 2)', '128', '10', '8%'],
            ['Highway',       '(4, 4, 1)', '16',  '4',  '25%'],
            ['LowerDeeps_A',  '(6, 6, 4)', '144', '6',  '4%'],
            ['OrcTown_A',     '(3, 3, 1)', '9',   '2',  '22%'],
            ['Elevator_B',    '(6, 6, 4)', '144', '0',  '0%'],
        ],
        col_widths=[1.8, 1.2, 0.8, 0.8, 0.8])
    doc.add_paragraph(
        'Z > 1 means a multi-level zone. Z=4 is a 4-storey mini-tower.')

    # --- NEW SECTION: Landmark connections ------------------------------
    add_heading(doc, 'Landmark connections — GuaranteedConnections', level=1)
    doc.add_paragraph(
        'A landmark\'s GuaranteedConnections[] is an array of GameplayTags '
        'in the form "World.Landmark.<row name>". Each entry names another '
        'landmark that the world generator MUST connect to this one. Only '
        'landmarks can connect via this mechanism — deck bubbles do not.')
    doc.add_paragraph(
        'Example: Chapter2.ElvenQuarterPromenade declares connections to '
        'Chapter2.MinesHeroShot and Chapter2.ElvenQuarterEntrance, so '
        'wherever the Promenade spawns, the generator wires a path to both '
        'peer landmarks.')
    doc.add_paragraph(
        'To guarantee a deck bubble appears near a landmark, put both in '
        'the same zone and set the deck entry to Appearances=Required.')

    # --- NEW SECTION: Appearances enum ----------------------------------
    add_heading(doc, 'Appearances enum — Required / Single / Multiple', level=1)
    add_table(doc,
        headers=['Value', 'Meaning'],
        rows=[
            ['Required', 'Guaranteed — at least one instance spawns in the zone'],
            ['Single',   'At most one (may or may not appear)'],
            ['Multiple', 'Zero or more (can repeat if generator chooses)'],
        ],
        col_widths=[1.2, 5.6])
    doc.add_paragraph(
        'Historical note: the Update 1 (March 2024) game had only 2 entries '
        'in TombsPassages (BB_DwarfHall1 Multiple + one other), which meant '
        'DwarfHall1 effectively always appeared in Tomb-zone passages. The '
        'modern game has a larger, more diverse passage pool, so the same '
        'Multiple flag now gets diluted. To restore "always present" '
        'behaviour on a modern deck, change Single/Multiple to Required.')

    # --- NEW SECTION: Build pipeline gotchas ----------------------------
    add_heading(doc, 'Build pipeline gotchas', level=1)

    add_heading(doc, '1. retoc to-zen output path must have .utoc extension',
                level=2)
    doc.add_paragraph(
        'The CLI\'s <OUTPUT> argument literally expects the .utoc filename, '
        'not a base name. Pass "SandboxMod_P.utoc". If you pass "SandboxMod_P" '
        '(no extension), retoc writes the .utoc at that exact path without '
        'the extension, and the IoStore runtime silently fails to load the '
        'container. The .pak and .ucas will exist but have no effect.')

    add_heading(doc, '2. Bundle only DataTables you actually changed', level=2)
    doc.add_paragraph(
        'Round-tripping an untouched DataTable through UAssetGUI (tojson then '
        'fromjson) can lose fidelity on complex nested types (Map properties '
        'on biomes, LinearColor arrays on zones, etc.). Observed symptom: '
        'in-game crash inside FAsyncLoadingThread with a bogus allocation '
        'size of 0xFFFFFFFE bytes — a signed-int underflow from a malformed '
        'name offset.')
    doc.add_paragraph(
        'The editor diffs each DataTable against its pristine .original.json '
        'sidecar and only bundles the ones that genuinely differ.')

    add_heading(doc, '3. Empty StructProperty arrays need DummyStruct', level=2)
    doc.add_paragraph(
        'UAssetGUI\'s JSON format emits a DummyStruct field on any '
        'ArrayProperty whose type is StructProperty and whose Value is '
        'empty. The DummyStruct preserves the inner struct type so fromjson '
        'can reconstruct the binary.')
    doc.add_paragraph(
        'If you empty out a struct array (e.g. remove the last landmark '
        'from LandmarkHandles) and lose the DummyStruct, fromjson throws: '
        '"Unable to reconstruct DummyStruct within empty StructProperty '
        'array". The editor preserves DummyStruct automatically: when you '
        'remove the last entry, the removed entry\'s shape becomes the new '
        'DummyStruct.')

    add_heading(doc, '4. UE4.27 UnrealPak cannot read Version 11 paks', level=2)
    doc.add_paragraph(
        'The Update 1 (March 2024) game pak uses pak Version 11. UE 4.27\'s '
        'UnrealPak errors with "PakEntry mismatch" on it. Use UE 5.7\'s '
        'UnrealPak instead:')
    add_mono(doc, 'C:/Program Files/Epic Games/UE_5.7/Engine/Binaries/Win64/UnrealPak.exe')

    add_heading(doc, '5. Sandbox save files bake the world at gen time', level=2)
    doc.add_paragraph(
        'Once a Sandbox Small run is generated, the layout is stored in the '
        'save file. Editing DT_Moria_Zones and deploying a new mod pak '
        'affects NEW worlds only. Existing saves keep their generated '
        'layout — start a fresh sandbox to see changes.')

    # --- NEW SECTION: Common confusions --------------------------------
    add_heading(doc, 'Common confusions', level=1)

    add_heading(doc, '"The Great Forge of Narvi" is NOT in any WorldGen DataTable',
                level=2)
    doc.add_paragraph(
        'It is a display name embedded inside the multi-origin bubble '
        'BD_BB_Chapter2_ElvenQuarterPromenade. The same physical room uasset '
        'has two labels — "The Elven Quarter" from one origin and "The '
        'Great Forge of Narvi" from another. It is not a separate zone, '
        'landmark, or deck entry. Editing bubble display names is out of '
        'scope for the WorldGen editor.')
    doc.add_paragraph(
        'Separately, the game has a Challenge Altar Blueprint '
        'BP_ChallengeAltar_Narvi referenced from DT_Moria_Waypoints row '
        'Durin_Muznakan2_Sandbox. Different Narvi — a progression altar, '
        'also out of scope today.')

    add_heading(doc, 'Campaign landmarks work fine in sandbox zones', level=2)
    doc.add_paragraph(
        'All 70 landmarks live in a single DT_Moria_Landmarks. There is no '
        'separate "sandbox landmark table". A sandbox zone can reference a '
        'Chapter2.* landmark directly by RowName — no duplication. The '
        'campaign landmark\'s GuaranteedConnections will not fire in sandbox '
        '(they point to other campaign landmarks that are not present), but '
        'the landmark\'s own bubble still spawns correctly.')

    add_heading(doc, 'Zones don\'t "own" landmarks — they reference them', level=2)
    doc.add_paragraph(
        'A zone\'s LandmarkHandles[] is a list of references. The same '
        'landmark can be referenced by multiple zones, but the world '
        'generator only places it once per run. If two zones claim the same '
        'transit landmark, one wins and the other\'s request is ignored.')

    # --- NEW SECTION: Editor UI features --------------------------------
    add_heading(doc, 'Editor UI features (SandboxZoneEditor.py)', level=1)

    add_heading(doc, 'Core tabs', level=2)
    add_table(doc,
        headers=['Tab', 'Purpose'],
        rows=[
            ['Zones', 'Edit 44 SandboxSmall zones (or all 91). Full per-field editing, editable landmark list, click-through to BubbleDeck/PassageDeck.'],
            ['Chapters', 'Edit all chapters. One-click "Add Book-Accurate Chapters 9-15".'],
            ['Biomes', 'View + edit DisplayName / Enabled. Object refs resolved to asset names.'],
            ['Bubbles', 'ZoneDeck DeckEntries editor. Zone-context strip shows both BubbleDeck and PassageDeck of the last zone you jumped from.'],
            ['Filters', 'Whitelist / Blacklist editor for DT_Moria_ZoneBubbleFilters.'],
            ['Landmarks', 'BaseBubbleName, Placement, PlayerStart, Challenge rating + GuaranteedConnections peer editor.'],
            ['Mappings', 'Zone → Chapter/Biome/Decks/Landmarks relationship view. Double-click a cell to jump to its editor tab.'],
            ['History', 'Row-centric diff vs pristine with field-level detail. Lock checkboxes protect changes from "Revert ALL" until unlocked.'],
            ['Map', '3D isometric canvas. Free yaw/pitch rotation (right-drag or Q/E/W/S keys). Per-chapter filter. Grid vs true-position layouts. Color-coded landmark markers (green star = player start, orange diamond = transit hub, cyan circle = set-piece).'],
        ],
        col_widths=[1.3, 5.5])

    add_heading(doc, 'Persistent state', level=2)
    doc.add_paragraph(
        'scripts/SandboxZoneEditor.ini stores:')
    add_table(doc,
        headers=['Section', 'Contents'],
        rows=[
            ['[sort]',  'Last-used column and direction per tab — sort state survives session restarts.'],
            ['[locks]', 'Protected changes marked "working and tested". Locked entries survive "Revert ALL".'],
        ],
        col_widths=[1.2, 5.6])

    add_heading(doc, 'Build flow', level=2)
    doc.add_paragraph(
        '1. Edit any DataTable. 2. Build > Preview build manifest — see '
        'exactly what will be bundled. 3. Build Mod Pak — produces '
        '~/Downloads/SandboxMod_v{VERSION}.zip containing the .pak/.ucas/'
        '.utoc triplet. 4. Extract to <Game>/Moria/Content/Paks/~mods/. '
        '5. Start a fresh sandbox to see the changes.')

    add_heading(doc, 'Revert flow', level=2)
    doc.add_paragraph(
        'History tab shows every change against pristine. Lock a change '
        '(checkbox) to mark it "working & tested" — survives Revert ALL. '
        'Revert at field / row / whole-table level. Individual revert on a '
        'locked change warns and requires unlocking first.')

    # --- NEW SECTION: Pristine baseline ---------------------------------
    add_heading(doc, 'Pristine baseline extraction', level=1)
    doc.add_paragraph(
        'The editor\'s revert and build-diff logic requires .original.json '
        'sidecars — a known-good pristine copy of each DataTable. Extract '
        'them once from the stock game pak:')
    add_mono(doc,
        '# 1. Unpack the stock game pak (UE 5.7 UnrealPak for version-11)\n'
        '"UnrealPak.exe" \\\n'
        '    "<Game>/Moria/Content/Paks/Moria-WindowsNoEditor.pak" \\\n'
        '    -Extract <scratch_dir> \\\n'
        '    "-Filter=Moria/Content/Tech/Data/GameWorld/DT_Moria*"\n\n'
        '# 2. Decompile each to JSON with UAssetGUI\n'
        'for name in DT_Moria_Zones DT_Moria_Chapters DT_Moria_Biomes \\\n'
        '           DT_Moria_ZoneDeck DT_Moria_ZoneBubbleFilters \\\n'
        '           DT_Moria_Landmarks; do\n'
        '    UAssetGUI.exe tojson \\\n'
        '        <scratch>/.../$name.uasset \\\n'
        '        experiments/worldgen_research/$name.original.json \\\n'
        '        VER_UE4_27\n'
        'done')
    doc.add_paragraph(
        'Sidecars are read-only reference data. Never written to by the '
        'editor.')

    doc.save(DOCX_PATH)
    print(f'Wrote {DOCX_PATH}')


def main():
    render_diagram()
    build_docx()


if __name__ == '__main__':
    main()
