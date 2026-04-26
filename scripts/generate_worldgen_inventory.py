"""
Generate a Word inventory of the Moria WorldGen DataTables.

Output: docs/Moria_WorldGen_Inventory.docx

Covers every row in:
  - DT_Moria_Chapters    (24 rows)
  - DT_Moria_Zones       (91 rows — campaign, sandbox, expedition)
  - DT_Moria_Biomes      (34 rows)
  - DT_Moria_Landmarks   (70 rows)
  - DT_Moria_ZoneDeck    (67 rows)
  - DT_Moria_ZoneBubbleFilters (12 rows)

For each entity, shows cross-references (which zones use which biome / landmark /
deck, which landmarks connect to which, etc.).
"""

from collections import defaultdict
from pathlib import Path
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
WGR = PROJECT_ROOT / 'experiments' / 'worldgen_research'
DOCS = PROJECT_ROOT / 'docs'
DOCS.mkdir(exist_ok=True)
OUT = DOCS / 'Moria_WorldGen_Inventory.docx'


# ---------- UAssetAPI JSON helpers ----------
def load_rows(fname):
    with open(WGR / fname, 'r', encoding='utf-8') as f:
        d = json.load(f)
    exports = d.get('Exports', [])
    return exports[0].get('Table', {}).get('Data', []) if exports else []


def prop(row, name):
    for p in row.get('Value', []) or []:
        if isinstance(p, dict) and p.get('Name') == name:
            return p
    return None


def get_enum(p):
    if p is None: return ''
    v = p.get('Value', '')
    return v.split('::', 1)[1] if isinstance(v, str) and '::' in v else str(v or '')


def get_rowname(p):
    if p is None: return ''
    for item in p.get('Value', []) or []:
        if isinstance(item, dict) and item.get('Name') == 'RowName':
            return str(item.get('Value', ''))
    return ''


def get_tagname(p):
    if p is None: return ''
    for item in p.get('Value', []) or []:
        if isinstance(item, dict) and item.get('Name') == 'TagName':
            return str(item.get('Value', ''))
    return ''


def get_intvec(p):
    if p is None: return (0, 0, 0)
    for item in p.get('Value', []) or []:
        if isinstance(item, dict):
            v = item.get('Value')
            if isinstance(v, dict) and 'X' in v:
                return (int(v['X']), int(v['Y']), int(v['Z']))
    return (0, 0, 0)


def get_scalar(p, default=0):
    return default if p is None else p.get('Value', default)


def get_landmarks(row):
    """List of (landmark_rowname, placement) tuples."""
    p = prop(row, 'LandmarkHandles')
    if p is None: return []
    out = []
    for entry in p.get('Value', []) or []:
        if not isinstance(entry, dict): continue
        lm = ''; placement = ''
        for sub in entry.get('Value', []) or []:
            if not isinstance(sub, dict): continue
            if sub.get('Name') == 'Landmark':
                lm = get_rowname(sub)
            elif sub.get('Name') == 'Placement':
                placement = get_enum(sub)
        if lm:
            out.append((lm, placement))
    return out


def get_deck_entries(row):
    p = prop(row, 'DeckEntries')
    if p is None: return []
    out = []
    for entry in p.get('Value', []) or []:
        if not isinstance(entry, dict): continue
        bubble = appearances = ''
        ze = False
        for sub in entry.get('Value', []) or []:
            if not isinstance(sub, dict): continue
            n = sub.get('Name')
            if n == 'Bubble':
                bubble = str(sub.get('Value', ''))
            elif n == 'Appearances':
                appearances = get_enum(sub)
            elif n == 'bZoneEntrance':
                ze = bool(sub.get('Value', False))
        out.append((bubble, appearances, ze))
    return out


def get_landmark_connections(row):
    p = prop(row, 'GuaranteedConnections')
    if p is None: return []
    out = []
    for item in p.get('Value', []) or []:
        if not isinstance(item, dict): continue
        inner = item.get('Value')
        tag = ''
        if isinstance(inner, str):
            tag = inner
        elif isinstance(inner, list):
            for sub in inner:
                if isinstance(sub, dict) and sub.get('Name') == 'TagName':
                    tag = str(sub.get('Value', ''))
                    break
        if tag.startswith('World.Landmark.'):
            out.append(tag[len('World.Landmark.'):])
        elif tag:
            out.append(tag)
    return out


def filter_list(row, name):
    p = prop(row, name)
    if p is None: return []
    out = []
    for item in p.get('Value', []) or []:
        if isinstance(item, dict):
            v = item.get('Value')
            if v: out.append(str(v))
    return out


# ---------- docx helpers ----------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], '2E5597')
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            c = table.rows[r_idx].cells[c_idx]
            c.text = str(val) if val is not None else ''
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(r.cells):
                    r.cells[i].width = Inches(w)
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return h


# ---------- main ----------

def build():
    chapters = load_rows('DT_Moria_Chapters.json')
    zones = load_rows('DT_Moria_Zones.json')
    biomes = load_rows('DT_Moria_Biomes.json')
    landmarks = load_rows('DT_Moria_Landmarks.json')
    decks = load_rows('DT_Moria_ZoneDeck.json')
    filters_ = load_rows('DT_Moria_ZoneBubbleFilters.json')

    # Cross-indices
    zones_by_chapter = defaultdict(list)
    zones_by_biome = defaultdict(list)
    zones_by_deck = defaultdict(list)
    zones_by_landmark = defaultdict(list)
    for zr in zones:
        ch = get_rowname(prop(zr, 'Chapter'))
        bi = get_tagname(prop(zr, 'Biome'))
        bd = get_rowname(prop(zr, 'BubbleDeck'))
        pd = get_rowname(prop(zr, 'PassageDeck'))
        zones_by_chapter[ch].append(zr['Name'])
        zones_by_biome[bi].append(zr['Name'])
        if bd: zones_by_deck[bd].append(zr['Name'])
        if pd and pd != bd: zones_by_deck[pd].append(zr['Name'])
        for lm, _pl in get_landmarks(zr):
            zones_by_landmark[lm].append(zr['Name'])

    # Reverse landmark connections
    conns_out = {}
    conns_in = defaultdict(list)
    for lr in landmarks:
        name = lr['Name']
        out = get_landmark_connections(lr)
        conns_out[name] = out
        for tgt in out:
            conns_in[tgt].append(name)

    # Chapter sort key
    def ch_key(c):
        try:
            return (0, int(c.rsplit('-', 1)[-1]))
        except Exception:
            return (1, c or '')

    # ---- Document ----
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)

    # Inventory tables are up to 9.1" wide, so force landscape letter with
    # tighter margins for ~9.5" usable width.
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # Cover
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Moria WorldGen Inventory')
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x2A, 0x44)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run('Every chapter, zone, biome, landmark, deck, and filter — '
                    'with cross-references (campaign + sandbox + expedition)')
    sr.italic = True; sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # Summary counts
    add_heading(doc, 'Contents at a glance', level=1)
    add_table(doc,
        ['DataTable', 'Count', 'What a row is'],
        [
            ['DT_Moria_Chapters', len(chapters),
             'A Z-layer bucket (Level 1–4 up, Deeps 1–4 down, plus DLC).'],
            ['DT_Moria_Zones', len(zones),
             'A cell of the map with chapter/biome/decks/landmarks/tuning.'],
            ['DT_Moria_Biomes', len(biomes),
             'A preset bundle: audio, deco, rock, atmospherics.'],
            ['DT_Moria_Landmarks', len(landmarks),
             'A named anchor point (BaseBubbleName + connections).'],
            ['DT_Moria_ZoneDeck', len(decks),
             'A pool of bubbles a zone can spawn.'],
            ['DT_Moria_ZoneBubbleFilters', len(filters_),
             'Whitelist/blacklist applied over bubble selection.'],
        ],
        col_widths=[2.2, 0.7, 3.7])

    # -----------------------------------------------------------
    # CHAPTERS
    # -----------------------------------------------------------
    add_heading(doc, '1. Chapters', level=1)
    doc.add_paragraph(
        'Each chapter is a Z-layer bucket with enemy scaling. Chapter row '
        'names are the keys zones use in their Chapter field.')
    rows = []
    for cr in sorted(chapters, key=lambda r: ch_key(r['Name'])):
        n = cr['Name']
        zs = get_enum(prop(cr, 'ZoneSet'))
        cid = get_scalar(prop(cr, 'ChapterID'), 0)
        disp = str(get_scalar(prop(cr, 'DisplayName'), ''))
        layer = get_scalar(prop(cr, 'Layer'), 0)
        scale = get_scalar(prop(cr, 'EnemyScalingLevel'), 0)
        minz = get_scalar(prop(cr, 'MinZ'), 0)
        maxz = get_scalar(prop(cr, 'MaxZ'), 0)
        primez = get_scalar(prop(cr, 'PrimeZ'), 0)
        enabled = get_enum(prop(cr, 'EnabledState')) or 'Live'
        nz = len(zones_by_chapter.get(n, []))
        rows.append([n, zs, cid, disp, layer, scale, f'{minz}/{maxz}/{primez}',
                     nz, enabled])
    add_table(doc,
        ['Chapter', 'ZoneSet', 'ID', 'DisplayName', 'Layer', 'Scale',
         'MinZ/MaxZ/PrimeZ', '#Zones', 'Enabled'],
        rows,
        col_widths=[1.8, 1.1, 0.4, 1.3, 0.5, 0.5, 1.0, 0.6, 0.6])

    # -----------------------------------------------------------
    # ZONES (grouped by chapter)
    # -----------------------------------------------------------
    add_heading(doc, '2. Zones by Chapter', level=1)
    doc.add_paragraph(
        'All 91 zones grouped by their Chapter field. Campaign (Moria-*) and '
        'DLC (Moria-DurinTower, Moria-DimrillDale, Moria-TradingPost) zones '
        'included. Columns: Zone, Biome, BubbleDeck, PassageDeck, '
        'Position, Size, Landmarks.')

    zones_by_name = {z['Name']: z for z in zones}
    # also gather zones with empty chapter
    empty_ch_zones = [z for z in zones if not get_rowname(prop(z, 'Chapter'))]

    chapters_in_use = sorted(zones_by_chapter.keys(), key=ch_key)
    for ch in chapters_in_use:
        if not ch:
            add_heading(doc, '(unassigned chapter)', level=2)
        else:
            add_heading(doc, ch, level=2)
        rows = []
        for zname in sorted(zones_by_chapter[ch]):
            zr = zones_by_name[zname]
            biome = get_tagname(prop(zr, 'Biome')).replace('World.Biome.', '')
            bd = get_rowname(prop(zr, 'BubbleDeck'))
            pd = get_rowname(prop(zr, 'PassageDeck'))
            px, py, pz = get_intvec(prop(zr, 'Position'))
            sx, sy, sz = get_intvec(prop(zr, 'TargetSize'))
            lms = ', '.join(lm for lm, _ in get_landmarks(zr)) or '—'
            enabled = get_enum(prop(zr, 'EnabledState')) or 'Live'
            rows.append([zname, biome, bd or '—', pd or '—',
                         f'({px},{py},{pz})', f'{sx}x{sy}x{sz}',
                         lms, enabled])
        add_table(doc,
            ['Zone', 'Biome', 'BubbleDeck', 'PassageDeck', 'Pos', 'Size',
             'Landmarks', 'Enb'],
            rows,
            col_widths=[1.8, 1.0, 1.2, 1.2, 0.7, 0.5, 2.4, 0.3])

    # -----------------------------------------------------------
    # BIOMES
    # -----------------------------------------------------------
    add_heading(doc, '3. Biomes', level=1)
    doc.add_paragraph(
        'All 34 biomes. "Used by" shows how many SandboxSmall + campaign zones '
        'reference the biome tag World.Biome.<Name>.')
    rows = []
    for br in biomes:
        n = br['Name']
        disp = str(get_scalar(prop(br, 'DisplayName'), ''))
        enabled = get_enum(prop(br, 'EnabledState')) or 'Live'
        # Biomes are referenced by zones via GameplayTag World.Biome.<Name>.
        tag_full = f'World.Biome.{n}'
        users = zones_by_biome.get(tag_full, [])
        rows.append([n, disp, len(users),
                     ', '.join(sorted(users)[:8]) +
                     (f' … (+{len(users) - 8})' if len(users) > 8 else '') or '—',
                     enabled])
    add_table(doc,
        ['Biome', 'DisplayName', '#Zones', 'Used by (sample)', 'Enb'],
        rows,
        col_widths=[1.5, 1.6, 0.6, 3.8, 0.3])

    # -----------------------------------------------------------
    # LANDMARKS
    # -----------------------------------------------------------
    add_heading(doc, '4. Landmarks', level=1)
    doc.add_paragraph(
        'All 70 landmarks with BaseBubbleName, Placement, which zones reference '
        'each, and GuaranteedConnections (→ outgoing, ← incoming).')
    rows = []
    for lr in sorted(landmarks, key=lambda r: r['Name']):
        n = lr['Name']
        bb = str(get_scalar(prop(lr, 'BaseBubbleName'), ''))
        pl = get_enum(prop(lr, 'Placement'))
        start = 'Yes' if get_scalar(prop(lr, 'bPlayerStartLocation'), False) else ''
        cr = get_scalar(prop(lr, 'ChallengeRating'), 0)
        enabled = get_enum(prop(lr, 'EnabledState')) or 'Live'
        zs = zones_by_landmark.get(n, [])
        zones_sample = ', '.join(sorted(zs)[:4]) + (
            f' (+{len(zs) - 4})' if len(zs) > 4 else '')
        out = conns_out.get(n, [])
        inc = conns_in.get(n, [])
        conn = ''
        if out: conn += '→ ' + ', '.join(sorted(out))
        if inc:
            if conn: conn += '   '
            conn += '← ' + ', '.join(sorted(inc))
        rows.append([n, bb, pl, zones_sample or '—',
                     conn or '—', start, cr, enabled])
    add_table(doc,
        ['Landmark', 'BaseBubbleName', 'Place', 'Used by zones (sample)',
         'Connections (→ out / ← in)', 'Start', 'Chal', 'Enb'],
        rows,
        col_widths=[1.6, 1.8, 0.55, 1.9, 2.3, 0.3, 0.3, 0.35])

    # -----------------------------------------------------------
    # ZONE DECKS
    # -----------------------------------------------------------
    add_heading(doc, '5. ZoneDecks (bubble pools)', level=1)
    doc.add_paragraph(
        'All 67 deck rows. Each deck is a collection of BB_* bubble names with '
        'Appearances rules. Zones reference decks via BubbleDeck / PassageDeck.')
    rows = []
    for dr in sorted(decks, key=lambda r: r['Name']):
        n = dr['Name']
        enabled = get_enum(prop(dr, 'EnabledState')) or 'Live'
        entries = get_deck_entries(dr)
        users = zones_by_deck.get(n, [])
        sample_bb = ', '.join(b for b, _, _ in entries[:4]) + (
            f' … (+{len(entries) - 4})' if len(entries) > 4 else '')
        rows.append([n, len(entries), sample_bb or '—',
                     ', '.join(sorted(users)[:4]) +
                     (f' (+{len(users) - 4})' if len(users) > 4 else '') or '—',
                     enabled])
    add_table(doc,
        ['ZoneDeck', '#', 'Sample bubbles', 'Used by zones (sample)', 'Enb'],
        rows,
        col_widths=[2.0, 0.4, 2.5, 2.4, 0.35])

    # -----------------------------------------------------------
    # FILTERS
    # -----------------------------------------------------------
    add_heading(doc, '6. ZoneBubbleFilters', level=1)
    doc.add_paragraph('All 12 filter rows with whitelist / blacklist contents.')
    rows = []
    for fr in sorted(filters_, key=lambda r: r['Name']):
        n = fr['Name']
        wl = filter_list(fr, 'Whitelist')
        bl = filter_list(fr, 'Blacklist')
        enabled = get_enum(prop(fr, 'EnabledState')) or 'Live'
        rows.append([n,
                     ', '.join(wl) or '—',
                     ', '.join(bl) or '—',
                     enabled])
    add_table(doc,
        ['Filter', 'Whitelist', 'Blacklist', 'Enb'],
        rows,
        col_widths=[1.6, 2.8, 2.8, 0.35])

    # -----------------------------------------------------------
    # APPENDIX: Reverse indexes for quick lookup
    # -----------------------------------------------------------
    add_heading(doc, '7. Reverse lookup — Biome → Zones', level=1)
    rows = []
    for tag, zs in sorted(zones_by_biome.items()):
        name = tag.replace('World.Biome.', '') if tag else '(no biome)'
        rows.append([name, len(zs), ', '.join(sorted(zs))])
    add_table(doc,
        ['Biome', '#', 'Zones'],
        rows,
        col_widths=[1.6, 0.5, 5.0])

    add_heading(doc, '8. Reverse lookup — Landmark → Zones', level=1)
    rows = []
    for lm, zs in sorted(zones_by_landmark.items()):
        rows.append([lm, len(zs), ', '.join(sorted(zs))])
    add_table(doc,
        ['Landmark', '#', 'Zones'],
        rows,
        col_widths=[2.0, 0.4, 5.0])

    add_heading(doc, '9. Reverse lookup — ZoneDeck → Zones', level=1)
    rows = []
    for d, zs in sorted(zones_by_deck.items()):
        rows.append([d, len(zs), ', '.join(sorted(zs))])
    add_table(doc,
        ['ZoneDeck', '#', 'Zones'],
        rows,
        col_widths=[2.0, 0.4, 5.0])

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    build()
